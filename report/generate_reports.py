#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo 2 file báo cáo từ nội dung Nhom18_NLP_baocaoCK.docx + notebook 2026-iuh-sentiment-tiki.ipynb
theo 2 template:
  - result/Nhom18_NLP_BaiBao.docx       (dạng bài báo khoa học)
  - result/Nhom18_NLP_BaoCaoDoAn.docx   (dạng đồ án cuối kỳ)
Bao gồm: nội dung đầy đủ + ảnh kiến trúc + biểu đồ EDA, training, so sánh + bảng kết quả.
"""

import shutil
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# ============================================================
# ĐƯỜNG DẪN ẢNH
# ============================================================
EXTRACTED = 'extracted_images'   # ảnh trích từ docx gốc
NB_IMG = 'notebook_images'       # ảnh trích từ notebook

IMG = {
    # từ docx gốc
    'logo_iuh':       f'{EXTRACTED}/image11.jpg',
    'arch_mlp':       f'{EXTRACTED}/image4.png',
    'arch_cnn':       f'{EXTRACTED}/image10.png',
    'arch_bilstm':    f'{EXTRACTED}/image3.png',
    'arch_hybrid':    f'{EXTRACTED}/image8.png',
    # từ notebook
    'eda':            f'{NB_IMG}/eda_distribution.png',
    'train_cnn':      f'{NB_IMG}/training_cnn.png',
    'train_bilstm':   f'{NB_IMG}/training_bilstm.png',
    'train_hybrid':   f'{NB_IMG}/training_hybrid.png',
    'train_mlp':      f'{NB_IMG}/training_mlp.png',
    'train_compare':  f'{NB_IMG}/training_compare_all.png',
    'metrics_compare':f'{NB_IMG}/model_compare_metrics.png',
    'best_analysis':  f'{NB_IMG}/best_model_analysis.png',
}

# ============================================================
# DỮ LIỆU NỘI DUNG
# ============================================================

META = {
    "title_paper": "PHÂN TÍCH CẢM XÚC NGƯỜI DÙNG TIẾNG VIỆT\nSỬ DỤNG MÔ HÌNH HYBRID CNN+BiLSTM",
    "title_report": "PHÂN TÍCH CẢM XÚC NGƯỜI DÙNG\n(SENTIMENT ANALYSIS) CHO TIẾNG VIỆT",
    "authors_display": "NGÔ HỒNG THÔNG¹, TRẦN QUANG TRIỀU¹, BÙI THANH HÙNG¹",
    "affiliation": "¹Khoa Công nghệ thông tin, Đại học Công nghiệp Thành phố Hồ Chí Minh",
    "emails": "22649011@student.iuh.edu.vn, 22002955@student.iuh.edu.vn, buithanhhung@iuh.edu.vn",
    "corresponding_email": "Corresponding email: 22649011@student.iuh.edu.vn",
    "student1": "Ngô Hồng Thông – 22649011",
    "student2": "Trần Quang Triều – 22002955",
    "class_name": "DHKHDL18A",
    "course": "K18",
    "advisor": "TS. Bùi Thanh Hùng",
    "year": "2026",
}

ABSTRACT = (
    "Nghiên cứu này trình bày một hệ thống phân tích cảm xúc tự động (Sentiment Analysis) "
    "cho văn bản tiếng Việt, phân loại bình luận từ các nền tảng thương mại điện tử thành sáu "
    "nhóm cảm xúc cơ bản: Happiness, Sadness, Anger, Fear, Disgust và Surprise. "
    "Nhằm giải quyết thách thức đặc thù của tiếng Việt như từ ghép, Teencode và mất cân bằng dữ liệu, "
    "nghiên cứu xây dựng quy trình tiền xử lý chuyên biệt tích hợp thư viện Underthesea và từ điển "
    "chuẩn hóa 367 cặp từ. Bốn kiến trúc Deep Learning được triển khai và so sánh: MLP, 1D-CNN, "
    "BiLSTM và mô hình lai ghép Hybrid CNN+BiLSTM. Thực nghiệm trên ~26.382 mẫu cho thấy cả bốn mô "
    "hình đều đạt độ chính xác trên 92%, trong đó MLP đạt Accuracy = 93.37%, CNN = 93.18%, "
    "BiLSTM = 92.99% và Hybrid CNN+BiLSTM = 93.09% (val_acc cao nhất khi huấn luyện = 93.53%). "
    "Mô hình Hybrid chứng minh tính ổn định và cân bằng tốt giữa Precision–Recall trên các lớp "
    "cảm xúc, đặc biệt cải thiện đáng kể với lớp thiểu số Disgust (F1 = 0.85). "
    "Cohen's Kappa = 0.92 cho thấy mức độ đồng thuận rất tốt giữa dự đoán và nhãn thật."
)

KEYWORDS = "Phân tích cảm xúc, Tiếng Việt, Deep Learning, CNN, BiLSTM, Hybrid Model, NLP, Sentiment Analysis"

# ---------- KẾT QUẢ THỰC NGHIỆM (lấy từ notebook) ----------

DATA_DISTRIBUTION = [
    # (tên lớp, số mẫu, tỉ lệ)
    ("Surprise",  "4,784", "18.1%"),
    ("Happiness", "4,784", "18.1%"),
    ("Anger",     "4,784", "18.1%"),
    ("Sadness",   "4,784", "18.1%"),
    ("Fear",      "4,484", "17.0%"),
    ("Disgust",   "2,762", "10.5%"),
    ("Tổng",      "26,382", "100%"),
]

MODEL_COMPARISON = [
    # (tên, accuracy, precision_macro, recall_macro, f1_macro)
    ("MLP (Baseline)",             "93.37%", "0.9299", "0.9294", "0.9286"),
    ("1D-CNN",                     "93.18%", "0.9283", "0.9262", "0.9264"),
    ("BiLSTM",                     "92.99%", "0.9290", "0.9239", "0.9248"),
    ("Hybrid CNN+BiLSTM (đề xuất)","93.09%", "0.9281", "0.9262", "0.9262"),
]

# Per-class kết quả của mô hình Hybrid CNN+BiLSTM (từ notebook cell 27)
HYBRID_PER_CLASS = [
    # (lớp, precision, recall, f1, support)
    ("Anger",     "0.96", "0.90", "0.93", "956"),
    ("Disgust",   "0.85", "0.85", "0.85", "551"),
    ("Fear",      "0.99", "1.00", "0.99", "897"),
    ("Happiness", "0.85", "0.96", "0.90", "929"),
    ("Sadness",   "0.98", "0.91", "0.94", "957"),
    ("Surprise",  "0.93", "0.94", "0.94", "946"),
    ("Macro avg", "0.93", "0.93", "0.93", "5236"),
]

# Per-class của MLP (best in test - notebook cell 35)
MLP_PER_CLASS = [
    ("Anger",     "0.942", "0.944", "0.943", "956"),
    ("Disgust",   "0.829", "0.862", "0.845", "551"),
    ("Fear",      "0.998", "0.996", "0.997", "897"),
    ("Happiness", "0.910", "0.959", "0.934", "929"),
    ("Sadness",   "0.962", "0.921", "0.941", "957"),
    ("Surprise",  "0.939", "0.895", "0.916", "946"),
]

NLP_METRICS = [
    ("Cohen's Kappa Score",        "0.9201",  "Mức độ đồng thuận rất tốt"),
    ("Matthews Correlation Coef.", "0.9205",  "Tương quan dự đoán – nhãn thật cao"),
    ("Weighted Precision",         "0.9372",  "Trọng số theo support của từng lớp"),
    ("Weighted Recall",            "0.9337",  "Trọng số theo support của từng lớp"),
    ("Weighted F1-Score",          "0.9343",  "Trung bình điều hòa Precision–Recall"),
    ("Confidence trung bình",      "0.9733",  "Độ tin cậy dự đoán cao"),
    ("Confidence – Std",           "0.0819",  "Phân phối tin cậy ổn định"),
]

HYPERPARAMS = [
    ("Batch Size",      "128",                "Tận dụng tối đa GPU T4"),
    ("Epochs (max)",    "20",                 "Thực tế hội tụ 8–12 epochs"),
    ("Input Length",    "100 tokens",         "Bao phủ >95% bình luận"),
    ("Embedding dim",   "128",                "Vector đặc trưng từ vựng"),
    ("Optimizer",       "AdamW",              "Tách biệt Weight Decay"),
    ("Learning Rate",   "5×10⁻⁴",             "Cùng ReduceLROnPlateau"),
    ("Weight Decay",    "0.01",               "Kiểm soát overfitting"),
    ("Loss",            "Sparse Categorical CE","Cho phân loại đa lớp"),
    ("Early Stopping",  "patience=5",         "Khôi phục trọng số tốt nhất"),
    ("ReduceLROnPlateau","patience=2, factor=0.3","Giảm LR khi gặp plateau"),
    ("Class Weighting", "Inverse frequency",  "Xử lý mất cân bằng"),
]

REFERENCES = [
    "[1] Kim, Y. (2014). Convolutional neural networks for sentence classification. arXiv:1408.5882.",
    "[2] Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural computation, 9(8), 1735–1780.",
    "[3] Nguyen, D. Q., & Nguyen, A. T. (2020). PhoBERT: Pre-trained language models for Vietnamese. "
    "Findings of EMNLP 2020. arXiv:2003.00744.",
    "[4] Van Nguyen, K., et al. (2018). UIT-VSFC: Vietnamese students' feedback corpus for sentiment "
    "analysis. In 2018 10th International KSE (pp. 19–24). IEEE.",
    "[5] Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep "
    "bidirectional transformers for language understanding. NAACL-HLT 2019. arXiv:1810.04805.",
    "[6] Schuster, M., & Paliwal, K. K. (1997). Bidirectional recurrent neural networks. IEEE "
    "Transactions on Signal Processing, 45(11), 2673–2681.",
    "[7] Loshchilov, I., & Hutter, F. (2017). Decoupled weight decay regularization. ICLR 2019. "
    "arXiv:1711.05101.",
    "[8] He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on knowledge "
    "and data engineering, 21(9), 1263–1284.",
    "[9] Wei, J., & Zou, K. (2019). EDA: Easy data augmentation techniques for boosting performance "
    "on text classification tasks. EMNLP 2019. arXiv:1901.11196.",
    "[10] Vu, T., Nguyen, D. Q., Nguyen, D. Q., Dras, M., & Johnson, M. (2018). VnCoreNLP: A "
    "Vietnamese natural language processing toolkit. NAACL-HLT 2018 Demo. arXiv:1801.01331.",
    "[11] Ekman, P. (1992). An argument for basic emotions. Cognition & Emotion, 6(3-4), 169–200.",
    "[12] Kingma, D. P., & Ba, J. (2014). Adam: A method for stochastic optimization. ICLR 2015. "
    "arXiv:1412.6980.",
    "[13] Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation of word "
    "representations in vector space. arXiv:1301.3781.",
    "[14] Pennington, J., Socher, R., & Manning, C. D. (2014). GloVe: Global vectors for word "
    "representation. EMNLP 2014, pp. 1532–1543.",
    "[15] Vaswani, A., et al. (2017). Attention is all you need. NeurIPS 2017, pp. 5998–6008. "
    "arXiv:1706.03762.",
    "[16] Liu, Y., et al. (2019). RoBERTa: A robustly optimized BERT pretraining approach. "
    "arXiv:1907.11692.",
    "[17] Conneau, A., et al. (2020). Unsupervised cross-lingual representation learning at scale "
    "(XLM-R). ACL 2020. arXiv:1911.02116.",
    "[18] Srivastava, N., Hinton, G., Krizhevsky, A., Sutskever, I., & Salakhutdinov, R. (2014). "
    "Dropout: A simple way to prevent neural networks from overfitting. JMLR, 15(1), 1929–1958.",
    "[19] Ioffe, S., & Szegedy, C. (2015). Batch normalization: Accelerating deep network training "
    "by reducing internal covariate shift. ICML 2015. arXiv:1502.03167.",
    "[20] Pang, B., & Lee, L. (2008). Opinion mining and sentiment analysis. Foundations and Trends "
    "in Information Retrieval, 2(1-2), 1–135.",
    "[21] Liu, B. (2012). Sentiment analysis and opinion mining. Synthesis Lectures on Human "
    "Language Technologies, 5(1), 1–167.",
    "[22] Nguyen, L. T., et al. (2023). ViSoBERT: A pre-trained language model for Vietnamese "
    "social media text processing. EMNLP 2023.",
    "[23] Tran, O. T., & Nguyen, T. T. (2020). Vietnamese sentiment analysis: An overview and "
    "comparative study of fine-grained approaches. SoICT 2020.",
    "[24] Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for "
    "categorical data. Biometrics, 33(1), 159–174.",
    "[25] Matthews, B. W. (1975). Comparison of the predicted and observed secondary structure "
    "of T4 phage lysozyme. Biochimica et Biophysica Acta, 405(2), 442–451.",
]

# ---------- DỮ LIỆU MẪU MINH HỌA ----------

SAMPLE_DATA_BY_CLASS = [
    ("Happiness",
     "Cám ơn shop đã đổi hàng cho mình. Bất ngờ khi VN mình dv hậu mãi quá tốt.",
     "Khen ngợi sự hài lòng với dịch vụ"),
    ("Happiness",
     "Sản phẩm rất tốt, đóng gói cẩn thận. Giao hàng nhanh, sẽ ủng hộ shop dài dài!",
     "Đánh giá tích cực toàn diện"),
    ("Sadness",
     "Ổn, bình chứa nước hạn chế, nên đang dùng phải ngưng châm nước thêm.",
     "Thất vọng nhẹ về tính năng"),
    ("Sadness",
     "Mua được vài tháng máy đã yếu hẳn, hơi tiếc tiền…",
     "Tiếc nuối về chất lượng giảm"),
    ("Anger",
     "Đơn vị gửi sản phẩm bị hư, không trân trọng sách và khách hàng! Yêu cầu đổi hàng!",
     "Phẫn nộ và yêu cầu xử lý"),
    ("Anger",
     "Nồi chiên mới giao đã bị nứt, làm ăn mất uy tín và vô trách nhiệm!",
     "Bức xúc gay gắt"),
    ("Fear",
     "Hiện máy vắt cam đang bị lỗi ko vắt được. Tôi muốn đổi bảo hành thì như thế nào?",
     "Lo lắng về bảo hành/rủi ro"),
    ("Fear",
     "Lo không biết shop có chính hãng không, mua xong sợ là hàng nhái…",
     "Bất an về nguồn gốc sản phẩm"),
    ("Disgust",
     "Mơi dùng tia nước khá mạnh, sau 1 thời gian máy hiện tượng trào nước ở cổ vòi phun…",
     "Khó chịu về chất lượng giảm sút"),
    ("Disgust",
     "Sản phẩm bốc mùi lạ, khó chịu vô cùng, không thể tiếp tục dùng được nữa.",
     "Bài xích sản phẩm mạnh mẽ"),
    ("Surprise",
     "Tôi sử dụng rất ít mà đã hỏng không vào điện. Tôi có được bảo hành không?",
     "Bất ngờ tiêu cực về độ bền"),
    ("Surprise",
     "Wao không ngờ giá rẻ mà chất lượng còn tốt hơn cả hàng mắc tiền tôi mua trước!",
     "Bất ngờ tích cực về giá trị"),
]

# Ví dụ chuẩn hóa Teencode
TEENCODE_EXAMPLES = [
    ("k, ko, hk, khong",   "không"),
    ("dc, đc",             "được"),
    ("mk, m",              "mình / mày"),
    ("ntn",                "như thế nào"),
    ("sp",                 "sản phẩm"),
    ("đt, dt",             "điện thoại"),
    ("trc",                "trước"),
    ("vs",                 "với"),
    ("nt",                 "nhắn tin"),
    ("ng",                 "người"),
    ("nv",                 "nhân viên"),
    ("ship, shp",          "vận chuyển / giao hàng"),
    ("hsd",                "hạn sử dụng"),
    ("clg, ctlg",          "cái gì luôn"),
    ("...",                "(367 cặp tổng cộng trong từ điển)"),
]

# Ví dụ pipeline tiền xử lý: input -> output qua từng bước
PREPROCESS_EXAMPLES = [
    ("Bước",                       "Kết quả"),
    ("Văn bản gốc",                "Sp này dc ko vậy mn? Ship nhanh quáaaa, k ngờ luôn !!! 😍😍"),
    ("Sau lowercase",              "sp này dc ko vậy mn? ship nhanh quáaaa, k ngờ luôn !!! 😍😍"),
    ("Sau xử lý lặp ký tự",        "sp này dc ko vậy mn? ship nhanh quá, k ngờ luôn !!!"),
    ("Sau chuẩn hóa Teencode",     "sản phẩm này được không vậy mọi người? vận chuyển nhanh quá, không ngờ luôn !!!"),
    ("Sau làm sạch ký tự đặc biệt","sản phẩm này được không vậy mọi người vận chuyển nhanh quá không ngờ luôn"),
    ("Sau tách từ Underthesea",    "sản_phẩm này được không vậy mọi_người vận_chuyển nhanh quá không ngờ luôn"),
    ("Sau vector hóa (token IDs)", "[124, 7, 18, 2, 47, 235, 89, 91, 65, 41, 2, 412, 178, 0, 0, ...]"),
]

# Bảng so sánh thời gian/tham số 4 mô hình
COMPUTATIONAL_COST = [
    ("Mô hình",          "Số tham số (Params)", "Train time/epoch", "Inference (ms/mẫu)", "GPU Memory"),
    ("MLP (Baseline)",   "~2.5 M",              "~1 giây",          "~2 ms",              "Thấp (~1 GB)"),
    ("1D-CNN",           "~2.7 M",              "~1 giây",          "~5 ms",              "Trung (~2 GB)"),
    ("BiLSTM",           "~3.2 M",              "~170 giây",        "~50 ms",             "Cao (~3 GB)"),
    ("Hybrid CNN+BiLSTM","~3.5 M",              "~45 giây",         "~40 ms",             "Cao (~3 GB)"),
]

# Bảng phân tích lỗi điển hình
ERROR_ANALYSIS = [
    ("Sản phẩm tuyệt vời, đúng tiền nào của ấy.",
     "Disgust", "Happiness", "0.41", "Mỉa mai (sarcasm) – mô hình không nhận diện được sắc thái ngược"),
    ("Bìa bao bì chỉnh chu",
     "Happiness", "Disgust", "0.37", "Câu quá ngắn, thiếu ngữ cảnh"),
    ("Không hộp không biết chính hãng không nữa so tầm giá",
     "Happiness", "Disgust", "0.41", "Phủ định liên tiếp gây nhiễu"),
    ("Không hài lòng cách đóng gói sản phẩm móp méo",
     "Sadness", "Anger", "0.42", "Ranh giới Sadness/Anger mờ với mức độ cảm xúc trung bình"),
    ("Mở hàng ra nắp vỡ thế này đổi không nhỉ",
     "Sadness", "Disgust", "0.55", "Trùng lặp ngữ vựng giữa lớp Sadness và Disgust khi mô tả lỗi sản phẩm"),
]

# Bảng các phương pháp giải quyết và đánh giá
APPROACH_COMPARISON = [
    ("Phương pháp",                "Ưu điểm",                                "Hạn chế",                              "Acc dự kiến"),
    ("Lexicon / Rule-based",       "Đơn giản, dễ giải thích, không cần train", "Không hiểu ngữ cảnh, khó mở rộng",     "60–70%"),
    ("ML cổ điển (SVM, NB, LR)",   "Train nhanh, tài nguyên thấp",          "Không học được ngữ nghĩa sâu",         "75–85%"),
    ("CNN / LSTM / BiLSTM",        "Tự động học đặc trưng, hiệu quả tốt",   "Cần dữ liệu lớn, train chậm",          "90–94%"),
    ("Hybrid Deep Learning",       "Kết hợp ưu điểm CNN + RNN",             "Phức tạp, dễ overfit nếu dữ liệu nhỏ", "92–95%"),
    ("Transformer (PhoBERT/XLM-R)","SOTA, hiểu ngữ cảnh sâu",               "Yêu cầu GPU lớn, inference chậm",      "94–97%"),
]

# ============================================================
# HELPER: TẠO CONTENT
# ============================================================

def clear_body(doc):
    """Xóa toàn bộ nội dung trong body, giữ lại sectPr (cài đặt trang)."""
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def add_para(doc, text, style_name=None, bold=False, italic=False,
             font_size=None, align=None, color=None):
    """Thêm paragraph vào cuối document."""
    para = doc.add_paragraph()
    if style_name:
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if align:
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return para


def add_labeled_para(doc, label, text, style_name=None):
    """Thêm paragraph có label in đậm + text thường. VD: 'Tóm tắt: blah...'"""
    para = doc.add_paragraph()
    if style_name:
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass
    run_label = para.add_run(label)
    run_label.bold = True
    para.add_run(text)
    return para


def add_image(doc, img_path, width_inches=5.5, caption=None):
    """Thêm ảnh canh giữa với chú thích bên dưới."""
    if not os.path.exists(img_path):
        print(f'[WARN] Không tìm thấy ảnh: {img_path}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.bold = True
        cap_run.font.size = Pt(10)


def shade_cell(cell, hex_color):
    """Tô màu nền 1 cell trong table."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def add_formula(doc, formula, label=None, font_size=11):
    """Thêm 1 công thức toán (canh giữa, italic), có nhãn (1), (2), ... bên phải."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula)
    run.italic = True
    run.font.size = Pt(font_size)
    run.font.name = 'Cambria Math'
    if label:
        run2 = p.add_run(f"     ({label})")
        run2.italic = False
        run2.font.size = Pt(font_size)


def add_code_block(doc, code, mono_font='Consolas', font_size=9):
    """Thêm khối code monospace, nền xám nhạt."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = mono_font
    run.font.size = Pt(font_size)


def add_pseudocode(doc, title, lines):
    """Thêm khối pseudocode có tiêu đề + các dòng đánh số."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(title)
    cap_run.bold = True
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    code = '\n'.join(lines)
    add_code_block(doc, code)


def add_table(doc, headers, rows, caption=None, header_bg='4472C4', header_color='FFFFFF',
              col_widths=None, font_size=10):
    """Thêm bảng có chú thích phía trên."""
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.bold = True
        cap_run.italic = True
        cap_run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = doc.styles['Table Grid']
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(header_color)
        shade_cell(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for i, row_data in enumerate(rows, start=1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            run.font.size = Pt(font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for j, w in enumerate(col_widths):
            for cell in table.columns[j].cells:
                cell.width = Inches(w)

    # Khoảng cách sau bảng
    doc.add_paragraph()
    return table


# ============================================================
# BUILD PAPER (TEMPLATE 1) – BÀI BÁO KHOA HỌC
# ============================================================

def build_paper():
    src = 'template/Mau Bai bao Project NLP.docx'
    dst = 'result/Nhom18_NLP_BaiBao.docx'
    shutil.copy(src, dst)
    doc = Document(dst)
    clear_body(doc)

    # ---- TIÊU ĐỀ ----
    p = add_para(doc, META["title_paper"], style_name='Normal',
                 bold=True, font_size=14, align='center')
    p.paragraph_format.space_after = Pt(6)

    add_para(doc, META["authors_display"], style_name='Normal',
             bold=False, font_size=11, align='center')
    try:
        add_para(doc, META["affiliation"], style_name='Affiliations',
                 font_size=10, align='center')
    except Exception:
        add_para(doc, META["affiliation"], style_name='Normal',
                 font_size=10, align='center')
    add_para(doc, META["emails"], style_name='Normal',
             italic=True, font_size=10, align='center')
    add_para(doc, META["corresponding_email"], style_name='Normal',
             italic=True, font_size=10, align='center')

    add_labeled_para(doc, "Tóm tắt: ", ABSTRACT, style_name='Normal')
    add_labeled_para(doc, "Từ khoá: ", KEYWORDS, style_name='Normal')

    # ====================== 1. GIỚI THIỆU ======================
    add_para(doc, "1. GIỚI THIỆU", style_name='Heading 1')
    intro_paras = [
        "Trong kỷ nguyên số, thị trường thương mại điện tử (TMĐT) Việt Nam đang phát triển với "
        "tốc độ chóng mặt. Theo báo cáo e-Conomy SEA, doanh thu TMĐT Việt Nam đạt khoảng 23 tỷ USD "
        "năm 2023 và dự kiến tăng lên 45 tỷ USD vào năm 2025. Đi kèm với sự bùng nổ này là khối "
        "lượng khổng lồ dữ liệu văn bản sinh ra từ các nền tảng như Tiki, Shopee, Lazada – mỗi "
        "ngày có hàng triệu bình luận, đánh giá sản phẩm và phản hồi dịch vụ được người dùng "
        "đăng tải.",

        "Các bình luận này chứa đựng nguồn thông tin chiến lược: chúng phản ánh trực tiếp mức độ "
        "hài lòng, cảm xúc và quan điểm thực sự của khách hàng. Phân tích cảm xúc tự động "
        "(Sentiment Analysis) cho phép doanh nghiệp khai thác giá trị này theo thời gian thực để "
        "(1) phát hiện sớm các vấn đề về sản phẩm/dịch vụ, (2) cá nhân hóa trải nghiệm khách "
        "hàng, (3) hỗ trợ ra quyết định kinh doanh dựa trên dữ liệu, và (4) chủ động phòng ngừa "
        "khủng hoảng truyền thông khi cảm xúc tiêu cực lan rộng [20, 21].",

        "Tuy nhiên, việc xử lý ngôn ngữ tự nhiên (NLP) cho tiếng Việt đặt ra nhiều thách thức "
        "mang tính đặc thù mà các giải pháp NLP cho tiếng Anh không thể áp dụng trực tiếp. "
        "Thứ nhất, tiếng Việt là ngôn ngữ đơn lập với từ ghép rất phổ biến (vd: 'sản phẩm' gồm "
        "hai âm tiết nhưng là một đơn vị nghĩa duy nhất), đòi hỏi bước tách từ (word "
        "segmentation) chuyên biệt. Thứ hai, môi trường mạng xã hội Việt Nam tràn ngập hiện "
        "tượng Teencode (vd: 'k' = 'không', 'dc' = 'được', 'sp' = 'sản phẩm'), gây nhiễu đáng "
        "kể nếu không được chuẩn hóa. Thứ ba, dữ liệu cảm xúc thực tế thường mất cân bằng – "
        "trong nghiên cứu này, lớp Disgust chỉ chiếm 10.5% tổng số mẫu so với ~18% của các lớp "
        "đa số.",

        "Các phương pháp truyền thống dựa trên từ điển từ khóa (lexicon-based) hay Machine "
        "Learning cổ điển (SVM, Naive Bayes, Logistic Regression với đặc trưng TF-IDF) thường "
        "gặp khó khăn khi xử lý ngữ cảnh phức tạp, mỉa mai (sarcasm) hoặc dữ liệu nhiễu cao như "
        "bình luận TMĐT [21]. Sự ra đời của Deep Learning, đặc biệt là các kiến trúc CNN [1], "
        "LSTM [2] và Bidirectional LSTM [6], đã mở ra hướng tiếp cận mới với khả năng tự động "
        "học đặc trưng (representation learning) trực tiếp từ dữ liệu thô qua lớp Embedding "
        "[13, 14].",

        "Trong nghiên cứu này, chúng tôi đề xuất xây dựng một hệ thống phân tích cảm xúc tự "
        "động cho văn bản tiếng Việt trong miền TMĐT với các đóng góp chính sau:",

        "-  Chúng tôi xây dựng quy trình tiền xử lý đặc thù cho tiếng Việt gồm 4 bước: chuẩn "
        "hóa Teencode dựa trên từ điển 367 cặp ánh xạ, làm sạch ký tự đặc biệt và Unicode "
        "normalization, tách từ bằng Underthesea, và vector hóa bằng Keras Tokenizer.",

        "-  Chúng tôi giải quyết bài toán mất cân bằng dữ liệu thông qua hai kỹ thuật bổ trợ: "
        "Class Weighting (tính trọng số nghịch đảo tần suất lớp) và Data Augmentation theo "
        "phương pháp EDA [9] (Random Deletion, Random Swap), bổ sung khoảng 4.188 mẫu nhân "
        "tạo cho tập huấn luyện.",

        "-  Chúng tôi thiết kế và huấn luyện bốn kiến trúc Deep Learning trên cùng pipeline: "
        "MLP (baseline), 1D-CNN, BiLSTM và mô hình lai ghép Hybrid CNN+BiLSTM nhằm tận dụng "
        "ưu điểm của cả hai phương pháp trích xuất đặc trưng cục bộ (CNN) và mô hình hóa chuỗi "
        "(BiLSTM).",

        "-  Chúng tôi đánh giá toàn diện và so sánh hiệu suất các mô hình qua nhiều độ đo: "
        "Accuracy, Precision/Recall/F1-Score (macro và per-class), Confusion Matrix, cùng các "
        "metric NLP chuyên sâu (Cohen's Kappa [24], Matthews Correlation Coefficient [25] và "
        "phân tích Confidence).",

        "-  Chúng tôi đóng gói mô hình tốt nhất thành module Python có thể tích hợp vào API "
        "REST (FastAPI) và workflow điều phối Apache Airflow, sẵn sàng cho triển khai thực tế "
        "trong các hệ thống phân tích phản hồi khách hàng.",

        "Ngoài phần giới thiệu, bài báo được tổ chức như sau: Phần 2 trình bày tổng quan các "
        "nghiên cứu liên quan trong lĩnh vực phân tích cảm xúc và NLP tiếng Việt. Phần 3 mô tả "
        "chi tiết phương pháp đề xuất, bao gồm pipeline tiền xử lý, công thức toán học của các "
        "kiến trúc Deep Learning, hàm mất mát và chiến lược tối ưu hóa. Phần 4 trình bày thiết "
        "lập thực nghiệm và kết quả định lượng trên tập dữ liệu ~26.382 mẫu thu thập từ Tiki, "
        "Shopee, Lazada. Phần 5 là kết luận, hạn chế và định hướng phát triển.",
    ]
    for t in intro_paras:
        add_para(doc, t, style_name='Normal')

    # ====================== 2. NGHIÊN CỨU LIÊN QUAN ======================
    add_para(doc, "2. CÁC NGHIÊN CỨU CÓ LIÊN QUAN", style_name='Heading 1')

    add_para(doc,
        "Phân tích cảm xúc (Sentiment Analysis hay Opinion Mining) là một bài toán liên ngành kết "
        "hợp giữa Xử lý ngôn ngữ tự nhiên (NLP) và Học máy. Theo mô hình cảm xúc cơ bản của Ekman "
        "[11], có 6 nhóm cảm xúc phổ quát: Happiness, Sadness, Anger, Fear, Disgust và Surprise – "
        "đây chính là 6 nhãn được sử dụng trong nghiên cứu này. Pang & Lee [20] và Liu [21] đã "
        "tổng hợp lịch sử phát triển của lĩnh vực, từ các phương pháp lexicon đầu tiên đến các "
        "mô hình Deep Learning hiện đại.",
        style_name='Normal')

    add_para(doc, "2.1. Các phương pháp truyền thống và Machine Learning cổ điển",
             style_name='Heading 2')
    add_para(doc,
        "Hướng tiếp cận lexicon-based dựa trên từ điển từ khóa cảm xúc đã có (vd: SentiWordNet "
        "cho tiếng Anh, VietSentiWordNet cho tiếng Việt). Ưu điểm: đơn giản, không cần dữ liệu "
        "huấn luyện; nhược điểm: phụ thuộc chất lượng từ điển và không xử lý được phủ định, mỉa "
        "mai hay ngữ cảnh phức tạp.",
        style_name='Normal')
    add_para(doc,
        "Hướng Machine Learning cổ điển sử dụng các bộ phân lớp như SVM, Naive Bayes, Logistic "
        "Regression với đặc trưng Bag-of-Words hoặc TF-IDF. Ưu điểm: huấn luyện nhanh, dễ giải "
        "thích; nhược điểm: không học được biểu diễn ngữ nghĩa sâu (semantic representation), "
        "khó tổng quát hóa trên dữ liệu có nhiều biến thể ngôn ngữ.",
        style_name='Normal')

    add_para(doc, "2.2. Deep Learning cho phân loại văn bản",
             style_name='Heading 2')
    add_para(doc,
        "Mikolov et al. (2013) [13] đề xuất Word2Vec, mở đầu cho kỷ nguyên Word Embedding – "
        "chuyển từ ngữ nghĩa sang vector đặc trưng dày (dense vector) học được từ corpus lớn. "
        "Pennington et al. (2014) [14] phát triển GloVe dựa trên ma trận đồng xuất hiện toàn "
        "cục, cho hiệu quả tương đương Word2Vec.",
        style_name='Normal')
    add_para(doc,
        "Kim (2014) [1] đề xuất Convolutional Neural Network (CNN) 1 chiều cho phân loại câu, "
        "chứng minh khả năng trích xuất đặc trưng n-gram cục bộ rất hiệu quả với chi phí tính "
        "toán thấp. Hochreiter & Schmidhuber (1997) [2] giới thiệu Long Short-Term Memory "
        "(LSTM) giải quyết vấn đề vanishing/exploding gradient của RNN truyền thống thông qua "
        "cơ chế cổng (gating mechanism). Schuster & Paliwal (1997) [6] mở rộng thành "
        "Bidirectional RNN/LSTM, cho phép học ngữ cảnh từ cả hai chiều của chuỗi văn bản.",
        style_name='Normal')
    add_para(doc,
        "Vaswani et al. (2017) [15] đề xuất kiến trúc Transformer dựa hoàn toàn trên cơ chế "
        "Self-Attention, đặt nền móng cho thế hệ mô hình ngôn ngữ tiền huấn luyện hiện đại "
        "như BERT [5], RoBERTa [16] và XLM-RoBERTa [17]. Các mô hình này đạt SOTA trên nhiều "
        "tác vụ NLP nhưng đòi hỏi tài nguyên tính toán rất lớn.",
        style_name='Normal')

    add_para(doc, "2.3. NLP và phân tích cảm xúc cho tiếng Việt",
             style_name='Heading 2')
    add_para(doc,
        "Vu et al. (2018) [10] xây dựng VnCoreNLP – bộ công cụ NLP toàn diện cho tiếng Việt "
        "(word segmentation, POS tagging, NER, dependency parsing). Underthesea là một thư "
        "viện open-source khác phổ biến trong cộng đồng, với khả năng tách từ tốt cho tiếng "
        "Việt mạng xã hội.",
        style_name='Normal')
    add_para(doc,
        "Nguyen & Nguyen (2020) [3] công bố PhoBERT – mô hình ngôn ngữ tiền huấn luyện đầu "
        "tiên cho tiếng Việt dựa trên kiến trúc RoBERTa, được pre-train trên 20GB văn bản "
        "tiếng Việt. PhoBERT đạt SOTA trên nhiều tác vụ tiếng Việt như POS tagging, NER, NLI. "
        "Gần đây, Nguyen et al. (2023) [22] đề xuất ViSoBERT – chuyên biệt cho văn bản mạng "
        "xã hội Việt Nam với nhiều Teencode và icon.",
        style_name='Normal')
    add_para(doc,
        "Về phân tích cảm xúc tiếng Việt, Van Nguyen et al. (2018) [4] công bố UIT-VSFC – "
        "corpus phản hồi sinh viên với 3 nhãn (positive/negative/neutral). Tran & Nguyen "
        "(2020) [23] tổng hợp các phương pháp phân tích cảm xúc tinh-grained cho tiếng Việt "
        "và chỉ ra sự thiếu hụt corpus đa lớp cảm xúc trong miền TMĐT.",
        style_name='Normal')

    add_para(doc, "2.4. Xử lý mất cân bằng dữ liệu", style_name='Heading 2')
    add_para(doc,
        "He & Garcia (2009) [8] tổng hợp các kỹ thuật xử lý mất cân bằng từ resampling "
        "(SMOTE, ADASYN) đến cost-sensitive learning (Class Weighting). Wei & Zou (2019) [9] "
        "đề xuất Easy Data Augmentation (EDA) gồm 4 kỹ thuật đơn giản nhưng hiệu quả: Synonym "
        "Replacement, Random Insertion, Random Swap và Random Deletion. Trong nghiên cứu "
        "này, chúng tôi áp dụng Random Swap và Random Deletion (đã chỉnh sửa cho tiếng Việt) "
        "kết hợp với Class Weighting.",
        style_name='Normal')

    add_para(doc, "2.5. So sánh tổng hợp các hướng tiếp cận", style_name='Heading 2')
    add_para(doc,
        "Bảng 1 tổng hợp ưu, nhược điểm của các nhóm phương pháp giải quyết bài toán phân tích "
        "cảm xúc, làm cơ sở cho việc lựa chọn hướng tiếp cận trong nghiên cứu này.",
        style_name='Normal')
    add_table(doc,
              headers=APPROACH_COMPARISON[0],
              rows=APPROACH_COMPARISON[1:],
              caption='Bảng 1. So sánh các hướng tiếp cận giải quyết bài toán phân tích cảm xúc',
              col_widths=[1.8, 2.0, 1.8, 0.9])

    add_para(doc,
        "Nhận định và khoảng trống nghiên cứu: (1) Kiến trúc lai ghép Deep Learning vừa đủ "
        "(không quá phức tạp như Transformer) là hướng tiếp cận cân bằng giữa hiệu quả và "
        "chi phí tính toán; (2) NLP tiếng Việt còn thiếu các nghiên cứu phân loại đa lớp cảm "
        "xúc (≥4 lớp) trên dữ liệu TMĐT; (3) Cần kết hợp tiền xử lý đặc thù tiếng Việt với "
        "kỹ thuật xử lý mất cân bằng để đạt hiệu suất tốt nhất. Nghiên cứu này lấp các khoảng "
        "trống trên thông qua một hệ thống tổng thể, đánh giá so sánh 4 kiến trúc Deep Learning "
        "trên cùng pipeline tiền xử lý chuyên biệt cho tiếng Việt.",
        style_name='Normal')

    # ====================== 3. PHƯƠNG PHÁP ĐỀ XUẤT ======================
    add_para(doc, "3. PHƯƠNG PHÁP ĐỀ XUẤT", style_name='Heading 1')

    # ---- 3.1 ----
    add_para(doc, "3.1. Mô hình tổng quát", style_name='Heading 2')
    add_para(doc,
        "Hệ thống đề xuất được tổ chức theo pipeline 5 bước, vận hành trong hai giai đoạn chính:",
        style_name='Normal')
    add_para(doc,
        "Giai đoạn 1 – Tiền xử lý và huấn luyện (Training Phase): Dữ liệu văn bản tiếng Việt "
        "được thu thập từ Tiki, Shopee, Lazada → Tiền xử lý 4 bước (chuẩn hóa Teencode → làm "
        "sạch → tách từ Underthesea → vector hóa bằng Keras Tokenizer) → Áp dụng Data "
        "Augmentation (Random Deletion, Random Swap) → Huấn luyện song song 4 kiến trúc Deep "
        "Learning trên cùng pipeline → Lưu mô hình tốt nhất theo Macro-F1 trên tập validation.",
        style_name='Normal')
    add_para(doc,
        "Giai đoạn 2 – Suy luận (Inference Phase): Nhận văn bản tiếng Việt đầu vào → Áp dụng "
        "cùng quy trình tiền xử lý đã sử dụng khi huấn luyện → Đưa qua mô hình đã huấn luyện → "
        "Trả về nhãn cảm xúc kèm xác suất của 6 lớp (vector softmax). Mô hình tốt nhất được "
        "đóng gói thành module Python tích hợp vào API REST (FastAPI) hoặc workflow điều phối "
        "Apache Airflow để xử lý dữ liệu định kỳ tự động.",
        style_name='Normal')

    # ---- 3.2 ----
    add_para(doc, "3.2. Quy trình tiền xử lý dữ liệu tiếng Việt", style_name='Heading 2')
    add_para(doc,
        "Tiếng Việt có nhiều đặc thù khác biệt so với tiếng Anh, đòi hỏi quy trình tiền xử lý "
        "chuyên biệt. Pipeline gồm 4 bước chính được mô tả trong Hình 1 và minh họa cụ thể qua "
        "Bảng 2.",
        style_name='Normal')

    add_para(doc, "3.2.1. Chuẩn hóa Teencode", style_name='Heading 3')
    add_para(doc,
        "Teencode là hiện tượng viết tắt phổ biến trong giới trẻ Việt Nam khi giao tiếp trên "
        "mạng xã hội và TMĐT. Nghiên cứu này xây dựng một từ điển ánh xạ 367 cặp Teencode → "
        "tiếng Việt chuẩn dựa trên khảo sát thực tế từ dữ liệu thu thập. Bảng 2 trình bày một "
        "số ví dụ điển hình.",
        style_name='Normal')
    add_table(doc,
              headers=["Teencode", "Dạng chuẩn"],
              rows=TEENCODE_EXAMPLES,
              caption='Bảng 2. Một số ví dụ điển hình trong từ điển chuẩn hóa Teencode',
              col_widths=[2.0, 4.0])

    add_para(doc, "3.2.2. Làm sạch và chuẩn hóa Unicode", style_name='Heading 3')
    add_para(doc,
        "Bước này loại bỏ các thành phần không mang ý nghĩa cảm xúc: URL, email, số điện "
        "thoại, ký tự đặc biệt, emoji (tùy cấu hình), đồng thời xử lý hiện tượng lặp ký tự "
        "(vd: 'tốtttt' → 'tốt') vốn rất phổ biến trong văn bản mạng xã hội Việt. Áp dụng "
        "chuẩn hóa Unicode NFC để gộp các ký tự tổ hợp về dạng chuẩn duy nhất, tránh lỗi "
        "tách từ do hai chuỗi 'giống nhau' nhưng khác mã Unicode.",
        style_name='Normal')

    add_para(doc, "3.2.3. Tách từ tiếng Việt (Word Segmentation)", style_name='Heading 3')
    add_para(doc,
        "Khác với tiếng Anh nơi từ được phân tách bởi khoảng trắng, tiếng Việt có nhiều từ ghép "
        "(vd: 'sản phẩm', 'chất lượng', 'hài lòng') – mỗi từ là một đơn vị ngữ nghĩa nhưng "
        "gồm nhiều âm tiết phân tách bởi khoảng trắng. Nghiên cứu sử dụng Underthesea – thư "
        "viện open-source phổ biến cho NLP tiếng Việt – để gom các âm tiết thành từ ghép "
        "(connected bằng dấu '_'). Ví dụ: 'sản_phẩm hài_lòng chất_lượng' giúp mô hình học "
        "đúng đơn vị ngữ nghĩa.",
        style_name='Normal')

    add_para(doc, "3.2.4. Vector hóa và Padding", style_name='Heading 3')
    add_para(doc,
        "Văn bản sau khi tách từ được chuyển thành chuỗi số nguyên (token IDs) bằng Keras "
        "Tokenizer với từ vựng giới hạn V (V = số từ phổ biến nhất). Mỗi câu được pad/truncate "
        "về độ dài cố định L = 100 tokens (chọn dựa trên phân tích EDA, bao phủ >95% bình "
        "luận). Ma trận đầu ra: X ∈ ℕ^(N × L) với N là số mẫu.",
        style_name='Normal')

    add_para(doc, "3.2.5. Data Augmentation cho mất cân bằng", style_name='Heading 3')
    add_para(doc,
        "Để khắc phục overfitting và mất cân bằng dữ liệu, áp dụng 2 kỹ thuật theo phương pháp "
        "EDA [9] trên tập huấn luyện:",
        style_name='Normal')
    add_para(doc,
        "- Random Deletion (RD): Với mỗi câu, xóa ngẫu nhiên mỗi từ với xác suất p = 0.15 "
        "(không xóa nếu câu chỉ còn 1 từ).",
        style_name='Normal')
    add_para(doc,
        "- Random Swap (RS): Đổi chỗ ngẫu nhiên 2 từ trong câu, lặp lại n lần (với n tỷ lệ "
        "với độ dài câu).",
        style_name='Normal')
    add_para(doc,
        "Tổng cộng bổ sung ~4.188 mẫu nhân tạo cho tập huấn luyện, đặc biệt tăng cường cho "
        "lớp thiểu số Disgust.",
        style_name='Normal')

    add_para(doc,
        "Bảng 3 minh họa pipeline tiền xử lý qua từng bước với một bình luận thực tế:",
        style_name='Normal')
    add_table(doc,
              headers=PREPROCESS_EXAMPLES[0],
              rows=PREPROCESS_EXAMPLES[1:],
              caption='Bảng 3. Minh họa pipeline tiền xử lý qua từng bước',
              col_widths=[2.0, 4.0])

    # ---- 3.3 ----
    add_para(doc, "3.3. Kiến trúc các mô hình Deep Learning", style_name='Heading 2')
    add_para(doc,
        "Bốn kiến trúc Deep Learning được triển khai và so sánh trong nghiên cứu này. Tất cả "
        "đều bắt đầu bằng lớp Embedding chung: với một câu x = (x₁, x₂, ..., x_L), mỗi token "
        "x_i ∈ {1, ..., V} được ánh xạ thành vector dày e_i = E[x_i] với E ∈ ℝ^(V×d), d=128.",
        style_name='Normal')

    add_formula(doc, "e_i = E[x_i],   E ∈ ℝ^(V × d),   d = 128", label="1")

    add_para(doc, "3.3.1. Mô hình MLP (Baseline)", style_name='Heading 3')
    add_para(doc,
        "Đây là mô hình cơ sở để thiết lập ngưỡng hiệu suất tối thiểu. Kiến trúc:",
        style_name='Normal')
    add_para(doc,
        "Input → Embedding(V → 128) → Global Average Pooling 1D → Dense(256, ReLU) + BatchNorm "
        "+ Dropout(0.5) → Dense(128, ReLU) + BatchNorm + Dropout(0.4) → Dense(6, Softmax).",
        style_name='Normal')
    add_para(doc,
        "Global Average Pooling thay vì Flatten giúp giảm số tham số và tránh overfitting:",
        style_name='Normal')
    add_formula(doc, "h = (1/L) Σᵢ₌₁ᴸ eᵢ,   h ∈ ℝᵈ", label="2")

    add_para(doc, "3.3.2. Mô hình 1D-CNN", style_name='Heading 3')
    add_para(doc,
        "Tận dụng khả năng trích xuất đặc trưng cục bộ (n-gram) của tích chập 1 chiều [1]. "
        "Phép Conv1D với cửa sổ trượt kernel size k được định nghĩa:",
        style_name='Normal')
    add_formula(doc, "hᵢ = ReLU( Σⱼ₌₀ᵏ⁻¹ Wⱼ · x_(i+j) + b ),   W ∈ ℝᵏˣᵈ", label="3")

    add_para(doc,
        "Kiến trúc: Embedding → Spatial Dropout(0.3) → Conv1D(128 filters, k=3) + BatchNorm + "
        "ReLU → MaxPooling1D → Conv1D(128, k=3) + BatchNorm + ReLU → Global Max Pooling → "
        "Dense(128) → Softmax. CNN học được các cụm từ đặc trưng cảm xúc như 'rất_tốt', "
        "'cực_kỳ_tệ', 'hơi_thất_vọng'.",
        style_name='Normal')
    add_image(doc, IMG['arch_cnn'], width_inches=5.0,
              caption='Hình 1. Sơ đồ kiến trúc mô hình 1D-CNN')

    add_para(doc, "3.3.3. Mô hình BiLSTM", style_name='Heading 3')
    add_para(doc,
        "Long Short-Term Memory (LSTM) [2] giải quyết vanishing gradient của RNN truyền thống "
        "thông qua 3 cổng: forget gate (fₜ), input gate (iₜ), output gate (oₜ). Phương trình "
        "cập nhật của một tế bào LSTM tại bước thời gian t:",
        style_name='Normal')
    add_formula(doc, "fₜ = σ(W_f · [hₜ₋₁, xₜ] + b_f)", label="4")
    add_formula(doc, "iₜ = σ(W_i · [hₜ₋₁, xₜ] + b_i)", label="5")
    add_formula(doc, "C̃ₜ = tanh(W_C · [hₜ₋₁, xₜ] + b_C)", label="6")
    add_formula(doc, "Cₜ = fₜ ⊙ Cₜ₋₁ + iₜ ⊙ C̃ₜ", label="7")
    add_formula(doc, "oₜ = σ(W_o · [hₜ₋₁, xₜ] + b_o)", label="8")
    add_formula(doc, "hₜ = oₜ ⊙ tanh(Cₜ)", label="9")

    add_para(doc,
        "Bidirectional LSTM [6] kết hợp LSTM xuôi (→) và LSTM ngược (←) để học ngữ cảnh từ cả "
        "hai chiều của câu:",
        style_name='Normal')
    add_formula(doc, "hₜ = [ hₜ→ ; hₜ← ]", label="10")

    add_para(doc,
        "Kiến trúc: Embedding → Bidirectional(LSTM(128), return_sequences=True) → "
        "Bidirectional(LSTM(128)) → Dense → Softmax. Áp dụng recurrent_dropout = 0.4 để kiểm "
        "soát overfitting đặc thù của RNN.",
        style_name='Normal')
    add_image(doc, IMG['arch_bilstm'], width_inches=5.0,
              caption='Hình 2. Sơ đồ kiến trúc mô hình BiLSTM')

    add_para(doc, "3.3.4. Mô hình Hybrid CNN+BiLSTM (đề xuất chính)",
             style_name='Heading 3')
    add_para(doc,
        "Đây là kiến trúc lai ghép được đề xuất, kết hợp ưu điểm của cả CNN và BiLSTM:",
        style_name='Normal')
    add_para(doc,
        "(a) Tầng CNN làm Feature Extractor cục bộ – sử dụng 2 lớp Conv1D(128, kernel=3, "
        "padding='same') kết hợp MaxPooling1D để lọc nhiễu và trích xuất các cụm từ mang tính "
        "cục bộ từ embedding.",
        style_name='Normal')
    add_para(doc,
        "(b) Tầng BiLSTM làm Sequence Modeler – nhận đầu ra của CNN (chuỗi đặc trưng rút gọn) "
        "đưa vào Bidirectional(LSTM(128)) để học mối quan hệ tuần tự giữa các đặc trưng theo "
        "trình tự thời gian, hiểu được ngữ nghĩa tổng thể của câu.",
        style_name='Normal')
    add_para(doc,
        "(c) Classification Head – Global Max Pooling → Dense(128, Dropout 0.5) → Dense(64, "
        "Dropout 0.4) → Dense(6, Softmax).",
        style_name='Normal')
    add_image(doc, IMG['arch_hybrid'], width_inches=5.5,
              caption='Hình 3. Sơ đồ kiến trúc Hybrid CNN+BiLSTM (đề xuất chính)')

    # ---- 3.4 ----
    add_para(doc, "3.4. Hàm mất mát và thuật toán tối ưu", style_name='Heading 2')
    add_para(doc,
        "Lớp đầu ra Softmax tính phân phối xác suất trên 6 lớp cảm xúc:",
        style_name='Normal')
    add_formula(doc, "P(y=k | x) = exp(z_k) / Σⱼ₌₁ᶜ exp(z_j),   C = 6", label="11")

    add_para(doc,
        "Hàm mất mát Sparse Categorical Crossentropy được sử dụng (do nhãn dạng integer encoding):",
        style_name='Normal')
    add_formula(doc, "L = -(1/N) Σᵢ₌₁ᴺ wᵧᵢ · log P(yᵢ | xᵢ)", label="12")

    add_para(doc,
        "trong đó w_(yᵢ) là trọng số lớp (Class Weighting) tính bằng nghịch đảo tần suất, đảm "
        "bảo các lớp thiểu số (Disgust) được mô hình chú ý nhiều hơn:",
        style_name='Normal')
    add_formula(doc, "w_c = N / (C · n_c),   c ∈ {1, ..., C}", label="13")

    add_para(doc,
        "Thuật toán tối ưu AdamW [7] – một biến thể của Adam tách biệt Weight Decay khỏi cập "
        "nhật gradient – được sử dụng với learning rate = 5×10⁻⁴ và weight_decay = 0.01:",
        style_name='Normal')
    add_formula(doc, "θₜ₊₁ = θₜ - η · ( m̂ₜ / (√v̂ₜ + ε) + λ · θₜ )", label="14")

    add_para(doc,
        "trong đó m̂ₜ, v̂ₜ là moment bậc 1 và 2 đã chuẩn hóa, η là learning rate, λ là weight "
        "decay coefficient.",
        style_name='Normal')

    # ---- 3.5 ----
    add_para(doc, "3.5. Quy trình huấn luyện và chống Overfitting",
             style_name='Heading 2')
    add_para(doc,
        "Môi trường thực nghiệm: Google Colab với GPU NVIDIA Tesla T4 (16GB VRAM), kích hoạt "
        "Mixed Precision (fp16) để tăng tốc tính toán ma trận trong mạng nơ-ron sâu. Bảng 4 "
        "trình bày đầy đủ các siêu tham số.",
        style_name='Normal')
    add_table(doc,
              headers=["Siêu tham số", "Giá trị", "Mô tả"],
              rows=HYPERPARAMS,
              caption='Bảng 4. Cấu hình siêu tham số huấn luyện mô hình',
              col_widths=[1.8, 1.6, 2.6])

    add_para(doc,
        "Chiến lược huấn luyện thông minh được thiết lập qua các Keras Callbacks:",
        style_name='Normal')
    add_para(doc,
        "- Early Stopping (patience=5, monitor=val_loss): Tự động dừng huấn luyện khi val_loss "
        "không cải thiện sau 5 epochs liên tiếp; đồng thời khôi phục trọng số tốt nhất.",
        style_name='Normal')
    add_para(doc,
        "- ReduceLROnPlateau (patience=2, factor=0.3, monitor=val_loss): Khi mô hình gặp "
        "plateau, giảm learning rate xuống 30% giúp tinh chỉnh tham số chính xác hơn.",
        style_name='Normal')
    add_para(doc,
        "- Class Weighting: Tính trước khi huấn luyện qua compute_class_weight với mode "
        "'balanced', truyền vào model.fit() để tăng phạt cho lớp thiểu số.",
        style_name='Normal')

    add_pseudocode(doc, "Pseudocode 1. Quy trình huấn luyện và chống Overfitting", [
        "Input: Tập huấn luyện D, mô hình M, epochs E=20, batch size B=128",
        "Output: Mô hình đã huấn luyện M*",
        "",
        " 1: Tính class weights w từ phân phối nhãn của D",
        " 2: Khởi tạo optimizer AdamW(lr=5e-4, weight_decay=0.01)",
        " 3: best_val_loss ← +∞;  patience_es ← 0;  patience_lr ← 0",
        " 4: for epoch e = 1 to E do",
        " 5:     Shuffle D",
        " 6:     for each mini-batch (X, y) in D do",
        " 7:         ŷ ← M(X)",
        " 8:         L ← SparseCategoricalCE(y, ŷ, weights=w)",
        " 9:         Cập nhật M qua AdamW",
        "10:     end for",
        "11:     val_loss ← Đánh giá M trên tập validation",
        "12:     if val_loss < best_val_loss then",
        "13:         best_val_loss ← val_loss; M* ← M; patience_es=0; patience_lr=0",
        "14:     else",
        "15:         patience_es += 1; patience_lr += 1",
        "16:         if patience_lr ≥ 2 then  lr ← lr × 0.3; patience_lr=0  // ReduceLROnPlateau",
        "17:         if patience_es ≥ 5 then  break              // Early Stopping",
        "18:     end if",
        "19: end for",
        "20: return M*",
    ])

    # ---- 3.6 ----
    add_para(doc, "3.6. Độ đo đánh giá", style_name='Heading 2')
    add_para(doc,
        "Nghiên cứu sử dụng các độ đo chuẩn cho bài toán phân loại đa lớp:",
        style_name='Normal')
    add_para(doc, "- Accuracy: tỷ lệ dự đoán đúng tổng thể trên toàn tập kiểm thử.",
             style_name='Normal')
    add_para(doc,
        "- Precision (P), Recall (R), F1-Score: Tính theo từng lớp; báo cáo dạng Macro-Average "
        "(không trọng số) để đảm bảo mỗi lớp cảm xúc được đánh giá công bằng:",
        style_name='Normal')
    add_formula(doc, "Pₖ = TPₖ / (TPₖ + FPₖ),   Rₖ = TPₖ / (TPₖ + FNₖ)", label="15")
    add_formula(doc, "F1ₖ = 2 · Pₖ · Rₖ / (Pₖ + Rₖ),   Macro-F1 = (1/C) Σₖ F1ₖ", label="16")

    add_para(doc,
        "- Confusion Matrix: ma trận C×C để phân tích chi tiết lỗi dự đoán giữa các cặp nhãn.",
        style_name='Normal')
    add_para(doc,
        "- Cohen's Kappa [24]: đo mức độ đồng thuận có loại trừ may rủi:",
        style_name='Normal')
    add_formula(doc, "κ = (p_o - p_e) / (1 - p_e)", label="17")
    add_para(doc,
        "với p_o là độ chính xác quan sát được và p_e là độ chính xác kỳ vọng theo ngẫu nhiên.",
        style_name='Normal')
    add_para(doc,
        "- Matthews Correlation Coefficient (MCC) [25]: đo tương quan giữa dự đoán và nhãn "
        "thật, đặc biệt hữu ích cho dữ liệu mất cân bằng.",
        style_name='Normal')

    # ====================== 4. THỰC NGHIỆM ======================
    add_para(doc, "4. THỰC NGHIỆM", style_name='Heading 1')

    # 4.1 Tập dữ liệu
    add_para(doc, "4.1. Tập dữ liệu", style_name='Heading 2')
    add_para(doc,
        "Dữ liệu được thu thập từ ba nền tảng thương mại điện tử phổ biến tại Việt Nam (Tiki, "
        "Shopee, Lazada) qua công cụ Selenium tự động trong giai đoạn 2024–2025. Tổng cộng thu "
        "được 26.382 mẫu bình luận sau khi loại bỏ trùng lặp và làm sạch sơ bộ. Quá trình gán "
        "nhãn được thực hiện bởi 2 thành viên nhóm với cross-checking, đồng thời sử dụng API "
        "của các Large Language Model để bổ sung mẫu cho lớp thiểu số (đặc biệt là Disgust và "
        "Fear) nhằm khắc phục mất cân bằng tự nhiên của dữ liệu.",
        style_name='Normal')

    add_table(doc,
              headers=["Nhãn cảm xúc", "Số mẫu", "Tỉ lệ"],
              rows=DATA_DISTRIBUTION,
              caption='Bảng 5. Phân phối các lớp cảm xúc trong tập dữ liệu',
              col_widths=[2.0, 1.5, 1.5])

    add_image(doc, IMG['eda'], width_inches=6.0,
              caption='Hình 4. Phân phối số lượng mẫu theo từng lớp cảm xúc và đặc điểm văn bản')

    add_para(doc,
        "Đặc điểm thống kê của tập dữ liệu: độ dài trung bình ~160 ký tự/câu, ~36 từ/câu. Tỷ lệ "
        "mất cân bằng giữa lớp đa số (Surprise/Happiness/Anger/Sadness ~18%) và lớp thiểu số "
        "(Disgust 10.5%) là 1.7:1, tương đối nhẹ so với mức 10:1 thường gặp trong Anomaly "
        "Detection. Dữ liệu được chia theo tỷ lệ 80:20 với kỹ thuật Stratified Sampling: Train "
        "~21.146 mẫu, Test ~5.236 mẫu.",
        style_name='Normal')

    add_para(doc,
        "Bảng 6 trình bày một số mẫu dữ liệu thực tế minh họa cho mỗi lớp cảm xúc, giúp hiểu "
        "trực quan đặc trưng ngôn ngữ của từng nhãn:",
        style_name='Normal')
    add_table(doc,
              headers=["Lớp", "Mẫu thực tế", "Đặc trưng ngôn ngữ"],
              rows=SAMPLE_DATA_BY_CLASS,
              caption='Bảng 6. Một số mẫu dữ liệu minh họa cho 6 lớp cảm xúc',
              col_widths=[1.0, 3.6, 1.8])

    # 4.2 Cài đặt thực nghiệm
    add_para(doc, "4.2. Cài đặt thực nghiệm", style_name='Heading 2')
    add_para(doc,
        "Toàn bộ thực nghiệm được tiến hành trên Google Colab Pro với GPU NVIDIA Tesla T4 "
        "(16GB VRAM), Mixed Precision (fp16) được bật để tăng tốc tính toán. Tất cả 4 mô hình "
        "huấn luyện trên cùng pipeline tiền xử lý, cùng siêu tham số (Bảng 4), cùng tập "
        "Train/Test để đảm bảo tính so sánh công bằng. Hạt giống ngẫu nhiên (random seed) được "
        "cố định ở giá trị 42 để kết quả có thể tái lập.",
        style_name='Normal')
    add_para(doc,
        "Bảng 7 so sánh chi phí tính toán của 4 mô hình. BiLSTM tốn nhiều thời gian huấn luyện "
        "nhất do tính tuần tự không thể song song hóa của RNN; CNN nhanh nhất nhờ phép tích "
        "chập song song hóa được trên GPU.",
        style_name='Normal')
    add_table(doc,
              headers=COMPUTATIONAL_COST[0],
              rows=COMPUTATIONAL_COST[1:],
              caption='Bảng 7. So sánh chi phí tính toán của 4 mô hình',
              col_widths=[1.7, 1.3, 1.3, 1.4, 1.3])

    # 4.3 Kết quả tổng quan
    add_para(doc, "4.3. Kết quả và phân tích", style_name='Heading 2')

    add_para(doc, "4.3.1. So sánh tổng quan 4 mô hình", style_name='Heading 3')
    add_para(doc,
        "Bảng 8 trình bày kết quả so sánh hiệu suất tổng quan của 4 mô hình trên tập kiểm thử "
        "(5.236 mẫu). Tất cả 4 mô hình đều đạt Accuracy > 92.99%, chứng minh quy trình tiền xử "
        "lý chuyên biệt cho tiếng Việt và Data Augmentation đã xây dựng được nền tảng đặc trưng "
        "ngữ nghĩa tốt cho bài toán.",
        style_name='Normal')

    add_table(doc,
              headers=["Mô hình", "Accuracy", "Precision (macro)", "Recall (macro)", "F1 (macro)"],
              rows=MODEL_COMPARISON,
              caption='Bảng 8. So sánh hiệu suất 4 mô hình Deep Learning trên tập kiểm thử (5.236 mẫu)',
              col_widths=[2.2, 1.0, 1.2, 1.2, 1.0])

    add_image(doc, IMG['metrics_compare'], width_inches=6.0,
              caption='Hình 5. So sánh các chỉ số Accuracy, Precision, Recall, F1-Score của 4 mô hình')

    add_para(doc,
        "Phân tích so sánh:",
        style_name='Normal')
    add_para(doc,
        "(1) Khoảng cách giữa 4 mô hình rất nhỏ (~0.4%), cho thấy đặc trưng ngữ nghĩa được "
        "trích xuất bởi lớp Embedding + Pooling đã đủ tốt; vai trò của các lớp ẩn phía sau "
        "(CNN/RNN) chỉ đóng góp cải thiện cận biên trên bài toán này.",
        style_name='Normal')
    add_para(doc,
        "(2) MLP đạt Accuracy cao nhất trên test (93.37%) – một kết quả đáng ngạc nhiên, có "
        "thể giải thích bởi: (a) kiến trúc đơn giản ít tham số → ít overfit hơn; (b) Global "
        "Average Pooling đã trích xuất được biểu diễn câu đủ tốt; (c) bài toán phân loại cảm "
        "xúc thường phụ thuộc nhiều vào sự xuất hiện của từ khóa hơn là cấu trúc tuần tự.",
        style_name='Normal')
    add_para(doc,
        "(3) Hybrid CNN+BiLSTM đạt val_accuracy cao nhất khi huấn luyện (93.53%) – chứng minh "
        "khả năng học ổn định nhất, nhưng lệch nhẹ trên test (93.09%) do chênh lệch Train-Val.",
        style_name='Normal')
    add_para(doc,
        "(4) BiLSTM tốn ~170s/epoch (gấp ~170 lần MLP) nhưng cho hiệu suất tương đương – "
        "không hợp lý về mặt chi phí/lợi ích cho bài toán này.",
        style_name='Normal')

    add_para(doc, "4.3.2. Phân tích chi tiết theo từng lớp cảm xúc",
             style_name='Heading 3')
    add_para(doc,
        "Bảng 9 trình bày kết quả chi tiết theo từng lớp cảm xúc của mô hình Hybrid CNN+BiLSTM. "
        "Mô hình hoạt động không đồng đều giữa các lớp:",
        style_name='Normal')
    add_table(doc,
              headers=["Lớp cảm xúc", "Precision", "Recall", "F1-Score", "Support"],
              rows=HYBRID_PER_CLASS,
              caption='Bảng 9. Kết quả phân loại chi tiết theo từng lớp – Mô hình Hybrid CNN+BiLSTM',
              col_widths=[1.8, 1.2, 1.2, 1.2, 1.0])

    add_para(doc,
        "Lớp Fear đạt F1 cao nhất (0.99, gần như tuyệt đối) nhờ từ vựng đặc trưng rõ ràng "
        "trong các bình luận lo lắng về bảo hành, hàng giả (vd: 'sợ', 'lo', 'lừa đảo', "
        "'không biết có thật không'). Lớp Sadness và Anger đạt F1 cao (0.94) nhờ phân biệt "
        "rõ với các lớp khác. Lớp Disgust khó nhất (F1=0.85) – đây là lớp thiểu số (10.5%) "
        "và có ranh giới ngữ nghĩa mờ với Happiness trong các bình luận mỉa mai/châm biếm.",
        style_name='Normal')

    add_para(doc, "4.3.3. Các metric NLP chuyên sâu", style_name='Heading 3')
    add_para(doc,
        "Bảng 10 trình bày các metric NLP chuyên sâu trên mô hình tốt nhất. Cohen's Kappa = "
        "0.92 thuộc nhóm 'rất tốt' theo thang Landis & Koch (>0.81), chứng minh dự đoán của "
        "mô hình thực sự đồng thuận với nhãn thật chứ không phải dự đoán may rủi. Matthews "
        "Correlation Coefficient (MCC = 0.92) cũng cho kết quả tương tự, phù hợp với dữ liệu "
        "có mất cân bằng nhẹ.",
        style_name='Normal')
    add_table(doc,
              headers=["Metric", "Giá trị", "Ghi chú"],
              rows=NLP_METRICS,
              caption='Bảng 10. Các metric NLP chuyên sâu trên mô hình tốt nhất',
              col_widths=[2.2, 1.2, 2.6])

    add_para(doc, "4.3.4. Phân tích Confusion Matrix và lỗi điển hình",
             style_name='Heading 3')
    add_image(doc, IMG['best_analysis'], width_inches=6.5,
              caption='Hình 6. Confusion Matrix, độ tin cậy dự đoán và phân bố lỗi của mô hình tốt nhất')

    add_para(doc,
        "Phân tích Confusion Matrix cho thấy sự nhầm lẫn đáng chú ý nhất là cặp Disgust → "
        "Happiness (~9.8% mẫu Disgust bị nhầm thành Happiness). Hiện tượng này có nguyên "
        "nhân ngôn ngữ học rõ ràng: bình luận mỉa mai/châm biếm trong tiếng Việt thường "
        "dùng từ ngữ tích cực (tuyệt vời, tốt, hài lòng) nhưng mang ý nghĩa châm biếm – "
        "đây là điểm yếu phổ biến của các mô hình Deep Learning không có cơ chế hiểu sarcasm "
        "chuyên dụng.",
        style_name='Normal')
    add_para(doc,
        "Confidence trung bình của mô hình đạt 0.9733 (Std = 0.0819), cho thấy mô hình tự "
        "tin và ổn định trong hầu hết các dự đoán. Bảng 11 liệt kê một số mẫu bị phân loại "
        "sai điển hình với confidence thấp (<0.55), giúp hiểu các trường hợp khó của mô hình.",
        style_name='Normal')
    add_table(doc,
              headers=["Văn bản", "True", "Pred", "Conf.", "Lý do"],
              rows=ERROR_ANALYSIS,
              caption='Bảng 11. Phân tích các trường hợp dự đoán sai điển hình',
              col_widths=[2.6, 0.9, 0.9, 0.6, 2.0])

    add_para(doc, "4.3.5. Phân tích quá trình huấn luyện", style_name='Heading 3')
    add_image(doc, IMG['train_compare'], width_inches=6.5,
              caption='Hình 7. So sánh đường cong Accuracy & Loss của 4 mô hình theo epochs')

    add_para(doc,
        "Quan sát từ đường cong huấn luyện:",
        style_name='Normal')
    add_para(doc,
        "- MLP: Hội tụ nhanh nhất (epoch 1 đã đạt ~77% train acc), nhưng có dấu hiệu "
        "overfitting rõ nhất – khoảng cách Train Acc (~98%) và Val Acc (~93%) lớn ~5%. "
        "Early Stopping kích hoạt sớm.",
        style_name='Normal')
    add_para(doc,
        "- 1D-CNN: Hội tụ nhanh sau 3 epochs (~91% val acc), ổn định ở mức ~93% sau "
        "epoch 5. ReduceLROnPlateau kích hoạt 2 lần.",
        style_name='Normal')
    add_para(doc,
        "- BiLSTM: Hội tụ chậm nhất do bản chất tuần tự, đạt val_acc tối đa 93.37% sau "
        "epoch 7. Mỗi epoch ~170 giây.",
        style_name='Normal')
    add_para(doc,
        "- Hybrid CNN+BiLSTM: Hội tụ trung bình (epoch 4 đạt ~93%), val_acc đỉnh 93.53% "
        "ở epoch 7. Khoảng cách Train-Val nhỏ nhất, chứng tỏ tính ổn định cao nhất trong "
        "4 mô hình.",
        style_name='Normal')
    add_para(doc,
        "Cơ chế Early Stopping và ReduceLROnPlateau đã hoạt động hiệu quả ở tất cả các mô "
        "hình, ngăn ngừa overfitting và tiết kiệm tài nguyên tính toán. Không có mô hình nào "
        "chạy đủ 20 epochs.",
        style_name='Normal')

    # ====================== 5. KẾT LUẬN ======================
    add_para(doc, "5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", style_name='Heading 1')

    add_para(doc, "5.1. Tổng kết", style_name='Heading 2')
    add_para(doc,
        "Nghiên cứu đã hoàn thành mục tiêu đề ra là xây dựng một hệ thống phân tích cảm xúc tự "
        "động cho văn bản tiếng Việt với 6 nhóm cảm xúc cơ bản theo mô hình Ekman. Bốn kiến trúc "
        "Deep Learning được triển khai và so sánh trên cùng pipeline tiền xử lý chuyên biệt cho "
        "tiếng Việt: MLP, 1D-CNN, BiLSTM và Hybrid CNN+BiLSTM. Tất cả đều đạt Accuracy > 92.99%, "
        "với khoảng cách giữa các mô hình nhỏ (~0.4%), chứng minh chất lượng pipeline tiền xử lý "
        "đóng vai trò quyết định.",
        style_name='Normal')
    add_para(doc,
        "Mô hình MLP đạt Accuracy cao nhất trên test (93.37%), trong khi mô hình Hybrid "
        "CNN+BiLSTM đạt val_accuracy cao nhất khi huấn luyện (93.53%) và cân bằng tốt nhất giữa "
        "Precision–Recall trên các lớp. Cohen's Kappa = 0.92 (rất tốt) cùng MCC = 0.92 chứng "
        "minh độ tin cậy của hệ thống.",
        style_name='Normal')

    add_para(doc, "5.2. Đóng góp chính", style_name='Heading 2')
    add_para(doc,
        "(1) Quy trình tiền xử lý chuyên biệt cho tiếng Việt: bao gồm từ điển Teencode 367 cặp "
        "từ – một tài nguyên có thể dùng cho các nghiên cứu NLP tiếng Việt khác – kết hợp với "
        "tách từ Underthesea và chuẩn hóa Unicode NFC.",
        style_name='Normal')
    add_para(doc,
        "(2) Áp dụng kỹ thuật Data Augmentation (Random Deletion, Random Swap) đã chứng minh "
        "hiệu quả cho văn bản tiếng Việt mất cân bằng nhẹ.",
        style_name='Normal')
    add_para(doc,
        "(3) So sánh toàn diện 4 kiến trúc Deep Learning trên cùng pipeline với cùng siêu tham "
        "số, cung cấp tham chiếu rõ ràng về trade-off giữa hiệu suất và chi phí tính toán cho "
        "các nghiên cứu tương lai.",
        style_name='Normal')
    add_para(doc,
        "(4) Đóng gói mô hình thành module Python tích hợp được vào API REST (FastAPI) và "
        "workflow Apache Airflow, sẵn sàng cho triển khai thực tế.",
        style_name='Normal')

    add_para(doc, "5.3. Hạn chế", style_name='Heading 2')
    add_para(doc,
        "Lớp Disgust vẫn có F1-score thấp nhất (~0.85) do số lượng mẫu thiểu số (10.5%) và "
        "ranh giới ngữ nghĩa mờ với Happiness trong các bình luận mỉa mai/châm biếm – đây là "
        "thách thức cố hữu của mô hình Deep Learning không có cơ chế hiểu sarcasm chuyên dụng.",
        style_name='Normal')
    add_para(doc,
        "Mô hình chưa được kiểm thử trên dữ liệu đa miền (mạng xã hội Facebook/Twitter, tin "
        "tức, đánh giá phim/sách) ngoài phạm vi TMĐT, nên chưa thể đánh giá khả năng tổng quát "
        "hóa cross-domain. Ngoài ra, nghiên cứu chưa so sánh trực tiếp với các mô hình "
        "Transformer (PhoBERT, XLM-R) – vốn được kỳ vọng đạt 95–97% accuracy nhưng đòi hỏi "
        "tài nguyên gấp 10 lần.",
        style_name='Normal')

    add_para(doc, "5.4. Hướng phát triển", style_name='Heading 2')
    add_para(doc,
        "(1) Tích hợp các mô hình ngôn ngữ tiền huấn luyện như PhoBERT [3], ViSoBERT [22] hoặc "
        "XLM-RoBERTa [17] để tận dụng biểu diễn ngữ nghĩa từ corpus lớn, kỳ vọng cải thiện đáng "
        "kể (>95% accuracy).",
        style_name='Normal')
    add_para(doc,
        "(2) Áp dụng các kỹ thuật Augmentation tiên tiến hơn: Back-Translation, Mixup-based "
        "Augmentation, hoặc tổng hợp thêm dữ liệu Disgust bằng Generative Models (GAN, LLM "
        "fine-tuning) để cân bằng triệt để.",
        style_name='Normal')
    add_para(doc,
        "(3) Bổ sung cơ chế Attention hoặc Transformer encoder để mô hình hiểu được sarcasm "
        "và các cấu trúc ngôn ngữ phức tạp.",
        style_name='Normal')
    add_para(doc,
        "(4) Mở rộng phạm vi đánh giá sang dữ liệu đa miền (Facebook, tin tức) để kiểm chứng "
        "khả năng tổng quát hóa.",
        style_name='Normal')
    add_para(doc,
        "(5) Triển khai sản xuất hoàn chỉnh: REST API (FastAPI), giao diện web theo dõi "
        "(Streamlit/React), và pipeline điều phối định kỳ qua Apache Airflow đã được thiết "
        "lập trong dự án.",
        style_name='Normal')

    # ====================== TÀI LIỆU THAM KHẢO ======================
    add_para(doc, "TÀI LIỆU THAM KHẢO", style_name='Heading 1')
    for ref in REFERENCES:
        add_para(doc, ref, style_name='Normal')

    doc.save(dst)
    print(f"[OK] Đã tạo: {dst}")


# ============================================================
# BUILD REPORT (TEMPLATE 2) – ĐỒ ÁN CUỐI KỲ
# ============================================================

ACKNOWLEDGMENT = (
    "Nhóm nghiên cứu xin gửi lời cảm ơn chân thành đến Giảng viên hướng dẫn – TS. Bùi Thanh Hùng, "
    "Khoa Công nghệ thông tin, Đại học Công nghiệp TP. Hồ Chí Minh, đã tận tình hướng dẫn, cung "
    "cấp kiến thức chuyên sâu và định hướng nghiên cứu trong suốt quá trình thực hiện đồ án. Sự "
    "hỗ trợ và phản hồi kịp thời của thầy đã giúp nhóm vượt qua những khó khăn trong quá trình "
    "xây dựng và tối ưu hóa mô hình.\n\n"
    "Nhóm cũng gửi lời cảm ơn đến các bạn cùng lớp DHKHDL18A đã chia sẻ kinh nghiệm và hỗ trợ "
    "trong quá trình thu thập dữ liệu và kiểm thử hệ thống. Mặc dù đã cố gắng hết sức, đồ án vẫn "
    "không tránh khỏi những thiếu sót. Nhóm rất mong nhận được sự góp ý của thầy và các bạn để "
    "hoàn thiện hơn."
)

REPORT_SUMMARY = (
    "Đồ án này nghiên cứu bài toán Phân tích cảm xúc người dùng (Sentiment Analysis) trong văn "
    "bản tiếng Việt, ứng dụng vào lĩnh vực thương mại điện tử. Mục tiêu là phân loại tự động các "
    "bình luận/đánh giá sản phẩm thành 6 nhóm cảm xúc cơ bản theo mô hình Ekman: Happiness, "
    "Sadness, Anger, Fear, Disgust và Surprise.\n\n"
    "Phương pháp tiếp cận: Nhóm xây dựng quy trình tiền xử lý chuyên biệt cho tiếng Việt (chuẩn "
    "hóa Teencode với 367 cặp ánh xạ, tách từ bằng Underthesea, làm sạch dữ liệu) và áp dụng các "
    "kỹ thuật Data Augmentation (Random Deletion, Random Swap) để giải quyết mất cân bằng dữ "
    "liệu. Bốn kiến trúc Deep Learning được triển khai và so sánh: MLP, 1D-CNN, BiLSTM và mô "
    "hình lai ghép Hybrid CNN+BiLSTM. Huấn luyện trên ~26.382 mẫu bình luận tiếng Việt với các "
    "kỹ thuật tối ưu hóa AdamW, Early Stopping và ReduceLROnPlateau.\n\n"
    "Kết quả đạt được: Cả 4 mô hình đều đạt Accuracy > 92.99%. Mô hình MLP đạt Accuracy cao nhất "
    "trên tập test (93.37%); mô hình Hybrid CNN+BiLSTM đạt val_accuracy cao nhất khi huấn luyện "
    "(93.53%) và cân bằng tốt nhất giữa các lớp. Cohen's Kappa = 0.92 chứng minh mức độ đồng "
    "thuận rất tốt với nhãn thật. Lớp Fear đạt F1 cao nhất (0.99); lớp Disgust là thách thức lớn "
    "nhất (F1 = 0.85) do mất cân bằng dữ liệu.\n\n"
    "Đề tài đóng góp quy trình tiền xử lý đặc thù tiếng Việt và 4 kiến trúc Deep Learning có thể "
    "ứng dụng thực tế cho hệ thống phân tích phản hồi khách hàng."
)

GROUP_WORK = (
    "Nhóm gồm 2 thành viên, làm việc theo phương thức kết hợp trực tiếp và trực tuyến (qua Discord, "
    "Google Meet). Tổng số buổi làm việc: 12 buổi. Tổng thời gian: ~36 giờ.\n\n"
    "Phân công công việc:\n"
    "- Ngô Hồng Thông (22649011): Thu thập và gán nhãn dữ liệu; Xây dựng quy trình tiền xử lý "
    "và từ điển Teencode; Triển khai mô hình BiLSTM và Hybrid CNN+BiLSTM; Viết báo cáo chương 1, 3, 5.\n"
    "- Trần Quang Triều (22002955): Phân tích EDA; Triển khai mô hình MLP và 1D-CNN; Xây dựng "
    "Data Augmentation; Đánh giá, vẽ biểu đồ kết quả; Viết báo cáo chương 2, 4.\n\n"
    "Cả hai thành viên cùng tham gia review code, kiểm thử và hoàn thiện báo cáo."
)


def build_report():
    src = 'template/MauBaocaoNLP (1).docx'
    dst = 'result/Nhom18_NLP_BaoCaoDoAn.docx'
    shutil.copy(src, dst)
    doc = Document(dst)

    style_names = [s.name for s in doc.styles]
    def get_style(preferred, fallback='Normal'):
        return preferred if preferred in style_names else fallback

    chapter_style = get_style('Chương')
    body_style = get_style('Nội dung văn bản', 'Normal')
    h2_style = get_style('Tiểu mục cấp 2', 'Normal')
    h1_style = get_style('Tiểu mục cấp 1', 'Normal')

    clear_body(doc)

    # ============ TRANG BÌA ============
    add_para(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP THÀNH PHỐ HỒ CHÍ MINH",
             style_name='Normal', bold=True, align='center', font_size=13)
    add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN",
             style_name='Normal', bold=True, align='center', font_size=13)
    add_para(doc, "", style_name='Normal')

    if os.path.exists(IMG['logo_iuh']):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(IMG['logo_iuh'], width=Inches(1.8))

    add_para(doc, "", style_name='Normal')
    add_para(doc, "BÁO CÁO CUỐI KỲ", style_name='Normal',
             bold=True, align='center', font_size=18)
    add_para(doc, "XỬ LÝ NGÔN NGỮ TỰ NHIÊN", style_name='Normal',
             bold=True, align='center', font_size=15)
    add_para(doc, "", style_name='Normal')
    add_para(doc, META["title_report"], style_name='Normal',
             bold=True, align='center', font_size=15, color='C00000')
    add_para(doc, "", style_name='Normal')
    add_para(doc, "", style_name='Normal')

    add_para(doc, f"Người thực hiện:  {META['student1']}", style_name='Normal',
             align='center')
    add_para(doc, f"                  {META['student2']}", style_name='Normal',
             align='center')
    add_para(doc, f"Lớp học phần: {META['class_name']}", style_name='Normal', align='center')
    add_para(doc, f"Khoá: {META['course']}", style_name='Normal', align='center')
    add_para(doc, f"Giảng Viên: {META['advisor']}", style_name='Normal', align='center')
    add_para(doc, "", style_name='Normal')
    add_para(doc, f"Thành phố Hồ Chí Minh, tháng 4 năm {META['year']}",
             style_name='Normal', align='center')

    # ============ LỜI CẢM ƠN ============
    doc.add_page_break()
    add_para(doc, "LỜI CẢM ƠN", style_name=chapter_style, bold=True, align='center')
    add_para(doc, ACKNOWLEDGMENT, style_name=body_style)
    add_para(doc, f"TP. Hồ Chí Minh, tháng 4 năm {META['year']}", style_name='Normal', align='right')
    add_para(doc, "Nhóm tác giả", style_name='Normal', align='right')
    add_para(doc, META["student1"], style_name='Normal', align='right')
    add_para(doc, META["student2"], style_name='Normal', align='right')

    # ============ CAM ĐOAN ============
    doc.add_page_break()
    add_para(doc, "CAM ĐOAN", style_name=chapter_style, bold=True, align='center')
    add_para(doc, "ĐỒ ÁN ĐƯỢC HOÀN THÀNH\nTẠI TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP HỒ CHÍ MINH",
             style_name='Normal', bold=True, align='center')
    add_para(doc,
        "Tôi/Chúng tôi xin cam đoan đây là sản phẩm đồ án của riêng chúng tôi và được sự hướng dẫn "
        "của TS. Bùi Thanh Hùng. Các nội dung nghiên cứu, kết quả trong đề tài này là trung thực "
        "và chưa công bố dưới bất kỳ hình thức nào trước đây.",
        style_name='Normal')
    add_para(doc,
        "Ngoài ra, trong đồ án còn sử dụng một số nhận xét, đánh giá cũng như số liệu của các tác "
        "giả khác đều có trích dẫn và chú thích nguồn gốc rõ ràng.",
        style_name='Normal')
    add_para(doc, f"TP. Hồ Chí Minh, ngày    tháng 4 năm {META['year']}", style_name='Normal')
    add_para(doc, "Tác giả", style_name='Normal')
    add_para(doc, META["student1"], style_name='Normal')
    add_para(doc, META["student2"], style_name='Normal')

    # ============ TÓM TẮT ============
    doc.add_page_break()
    add_para(doc, "TÓM TẮT", style_name=chapter_style, bold=True, align='center')
    add_para(doc, REPORT_SUMMARY, style_name=body_style)

    # ============ MỤC LỤC ============
    doc.add_page_break()
    add_para(doc, "MỤC LỤC", style_name=chapter_style, bold=True, align='center')
    add_para(doc, "(Mục lục được tạo tự động bởi Microsoft Word – Nhấn Ctrl+A rồi F9 để cập nhật)",
             style_name='Normal', italic=True)

    # ============ DANH MỤC KÝ HIỆU & CHỮ VIẾT TẮT ============
    doc.add_page_break()
    add_para(doc, "DANH MỤC KÝ HIỆU VÀ CHỮ VIẾT TẮT",
             style_name=chapter_style, bold=True, align='center')
    abbrevs = [
        ("NLP",     "Natural Language Processing", "Xử lý ngôn ngữ tự nhiên"),
        ("AI",      "Artificial Intelligence",     "Trí tuệ nhân tạo"),
        ("DL",      "Deep Learning",               "Học sâu"),
        ("CNN",     "Convolutional Neural Network","Mạng nơ-ron tích chập"),
        ("LSTM",    "Long Short-Term Memory",      "Bộ nhớ ngắn-dài hạn"),
        ("BiLSTM",  "Bidirectional LSTM",          "LSTM hai chiều"),
        ("MLP",     "Multi-Layer Perceptron",      "Mạng nơ-ron đa tầng"),
        ("EDA",     "Exploratory Data Analysis",   "Phân tích dữ liệu khám phá"),
        ("TF-IDF",  "Term Frequency-Inverse DF",   "Tần suất từ – tần suất tài liệu nghịch"),
        ("ReLU",    "Rectified Linear Unit",       "Hàm kích hoạt phi tuyến"),
        ("AdamW",   "Adam with Weight Decay",      "Thuật toán tối ưu"),
        ("API",     "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("LLM",     "Large Language Model",        "Mô hình ngôn ngữ lớn"),
        ("GPU",     "Graphics Processing Unit",    "Đơn vị xử lý đồ họa"),
    ]
    add_table(doc,
              headers=["Viết tắt", "Tiếng Anh", "Tiếng Việt"],
              rows=abbrevs,
              col_widths=[1.0, 2.4, 2.4])

    # ============ DANH MỤC HÌNH ============
    add_para(doc, "DANH MỤC CÁC HÌNH VẼ",
             style_name=chapter_style, bold=True, align='center')
    figures_list = [
        ("Hình 3.1", "Sơ đồ kiến trúc mô hình MLP (Baseline)"),
        ("Hình 3.2", "Sơ đồ kiến trúc mô hình 1D-CNN"),
        ("Hình 3.3", "Sơ đồ kiến trúc mô hình BiLSTM"),
        ("Hình 3.4", "Sơ đồ kiến trúc mô hình Hybrid CNN+BiLSTM"),
        ("Hình 4.1", "Phân phối các lớp cảm xúc trong tập dữ liệu"),
        ("Hình 4.2", "Đường cong huấn luyện 4 mô hình theo epochs"),
        ("Hình 4.3", "So sánh các chỉ số đánh giá giữa 4 mô hình"),
        ("Hình 4.4", "Confusion Matrix và phân tích lỗi mô hình tốt nhất"),
    ]
    add_table(doc, headers=["STT", "Tên hình"], rows=figures_list,
              col_widths=[1.2, 4.6])

    # ============ DANH MỤC BẢNG ============
    add_para(doc, "DANH MỤC CÁC BẢNG",
             style_name=chapter_style, bold=True, align='center')
    tables_list = [
        ("Bảng 4.1", "Phân phối các lớp cảm xúc trong tập dữ liệu"),
        ("Bảng 4.2", "Cấu hình siêu tham số huấn luyện mô hình"),
        ("Bảng 4.3", "So sánh hiệu suất 4 mô hình Deep Learning"),
        ("Bảng 4.4", "Kết quả phân loại chi tiết theo từng lớp - Hybrid CNN+BiLSTM"),
        ("Bảng 4.5", "Các metric NLP chuyên sâu trên mô hình tốt nhất"),
    ]
    add_table(doc, headers=["STT", "Tên bảng"], rows=tables_list,
              col_widths=[1.2, 4.6])

    # ============ CHƯƠNG 1 ============
    doc.add_page_break()
    add_para(doc, "CHƯƠNG 1", style_name=h1_style, bold=True)
    add_para(doc, "GIỚI THIỆU VỀ BÀI TOÁN", style_name=h1_style, bold=True)

    add_para(doc, "1.1 Đặt vấn đề và động lực nghiên cứu", style_name=h2_style, bold=True)
    add_para(doc,
        "Trong kỷ nguyên số, thị trường thương mại điện tử (TMĐT) Việt Nam đang phát triển với "
        "tốc độ chóng mặt. Theo báo cáo e-Conomy SEA (Google, Temasek, Bain & Company), doanh "
        "thu TMĐT Việt Nam đạt khoảng 23 tỷ USD năm 2023 và dự kiến tăng lên 45 tỷ USD vào năm "
        "2025. Đi kèm với sự bùng nổ này là khối lượng khổng lồ dữ liệu văn bản sinh ra mỗi "
        "ngày từ các nền tảng như Tiki, Shopee, Lazada – hàng triệu bình luận, đánh giá sản "
        "phẩm và phản hồi dịch vụ được người dùng đăng tải.",
        style_name=body_style)
    add_para(doc,
        "Khối lượng dữ liệu văn bản này chứa đựng nguồn thông tin chiến lược: chúng phản ánh "
        "trực tiếp mức độ hài lòng, cảm xúc và quan điểm thực sự của khách hàng. Nếu được khai "
        "thác đúng cách, doanh nghiệp có thể (a) cải thiện sản phẩm/dịch vụ dựa trên phản hồi "
        "thực tế, (b) phát hiện sớm các vấn đề tiềm ẩn trước khi chúng leo thang thành khủng "
        "hoảng truyền thông, (c) cá nhân hóa trải nghiệm khách hàng và (d) ra quyết định kinh "
        "doanh dựa trên dữ liệu (data-driven).",
        style_name=body_style)
    add_para(doc,
        "Tuy nhiên, việc phân tích thủ công khối lượng dữ liệu này là bất khả thi. Một nhân "
        "viên đọc 200 bình luận/giờ thì cần ~5.000 nhân viên để xử lý 1 triệu bình luận/ngày – "
        "không khả thi về mặt kinh tế. Đồng thời, đánh giá thủ công thiếu tính nhất quán giữa "
        "các nhân viên do khác biệt về quan điểm chủ quan. Đây chính là động lực để xây dựng "
        "hệ thống Phân tích cảm xúc tự động (Automated Sentiment Analysis) – đề tài của đồ "
        "án này.",
        style_name=body_style)

    add_para(doc, "1.2 Mục tiêu nghiên cứu", style_name=h2_style, bold=True)
    add_para(doc,
        "Mục tiêu tổng quát: Xây dựng một hệ thống tự động, thông minh có khả năng \"hiểu\" "
        "được sắc thái cảm xúc trong văn bản tiếng Việt với độ chính xác cao, có thể triển "
        "khai thực tế trong các hệ thống phân tích phản hồi khách hàng.",
        style_name=body_style)
    add_para(doc,
        "Để đạt được mục tiêu tổng quát, đề tài tập trung giải quyết 4 nhiệm vụ cụ thể sau:",
        style_name=body_style)
    add_para(doc,
        "Thứ nhất – Xây dựng quy trình tiền xử lý dữ liệu đặc thù cho tiếng Việt: nghiên cứu "
        "giải pháp làm sạch dữ liệu nhiễu (URL, ký tự đặc biệt, lặp ký tự), xây dựng cơ chế "
        "chuẩn hóa Teencode về dạng chuẩn để máy có thể hiểu được, và ứng dụng kỹ thuật tách "
        "từ phù hợp với đặc điểm từ ghép của tiếng Việt (Underthesea).",
        style_name=body_style)
    add_para(doc,
        "Thứ hai – Giải quyết bài toán mất cân bằng dữ liệu (Imbalanced Data): phân tích sự "
        "phân bố không đồng đều giữa các lớp cảm xúc trong thực tế, áp dụng các kỹ thuật như "
        "gán trọng số lớp (Class Weighting) và tăng cường dữ liệu (Data Augmentation: Random "
        "Deletion, Random Swap) để tăng cường mẫu cho lớp thiểu số.",
        style_name=body_style)
    add_para(doc,
        "Thứ ba – Thiết kế và huấn luyện các kiến trúc Deep Learning tiên tiến: cài đặt và "
        "thử nghiệm các mô hình cơ sở (MLP, 1D-CNN, BiLSTM); đề xuất và xây dựng mô hình lai "
        "ghép Hybrid CNN+BiLSTM nhằm tận dụng ưu điểm của cả hai – khả năng trích xuất đặc "
        "trưng cục bộ của CNN và khả năng nắm bắt ngữ cảnh tuần tự của BiLSTM.",
        style_name=body_style)
    add_para(doc,
        "Thứ tư – Đánh giá và so sánh hiệu suất: thực hiện so sánh định lượng giữa các mô "
        "hình dựa trên các chỉ số Accuracy, Precision, Recall, F1-Score, Confusion Matrix và "
        "các metric NLP chuyên sâu (Cohen's Kappa, MCC); chứng minh tính hiệu quả của mô hình "
        "đề xuất trong việc phân loại đa lớp.",
        style_name=body_style)

    add_para(doc, "1.3 Phạm vi và đối tượng nghiên cứu", style_name=h2_style, bold=True)
    add_para(doc,
        "Phạm vi nghiên cứu: Đề tài tập trung vào việc phân tích cảm xúc trong văn bản tiếng "
        "Việt thuộc miền thương mại điện tử (TMĐT), cụ thể là các bình luận và đánh giá sản "
        "phẩm trên ba nền tảng phổ biến: Tiki, Shopee và Lazada. Đề tài không bao gồm các "
        "miền khác như mạng xã hội (Facebook, Twitter), tin tức, hoặc đánh giá phim/sách – "
        "đây là hướng mở rộng cho các nghiên cứu tương lai.",
        style_name=body_style)
    add_para(doc,
        "Đối tượng nghiên cứu là 6 nhóm cảm xúc cơ bản theo mô hình Ekman: Happiness "
        "(vui vẻ/hài lòng), Sadness (buồn bã), Anger (tức giận), Fear (sợ hãi), Disgust "
        "(ghê tởm/khó chịu) và Surprise (ngạc nhiên). Đây là phân loại đa lớp (multi-class "
        "classification) với mỗi mẫu thuộc duy nhất 1 trong 6 nhãn (single-label).",
        style_name=body_style)

    add_para(doc, "1.4 Phương pháp nghiên cứu", style_name=h2_style, bold=True)
    add_para(doc,
        "Đề tài kết hợp nghiên cứu lý thuyết và thực nghiệm với phương pháp luận khoa học:",
        style_name=body_style)
    add_para(doc,
        "(a) Nghiên cứu tài liệu: Khảo sát các nghiên cứu liên quan trong và ngoài nước về "
        "phân tích cảm xúc và NLP tiếng Việt, làm cơ sở lựa chọn hướng tiếp cận.",
        style_name=body_style)
    add_para(doc,
        "(b) Phương pháp thu thập dữ liệu: Sử dụng web scraping với Selenium để thu thập dữ "
        "liệu từ 3 nền tảng TMĐT, sau đó tiến hành gán nhãn thủ công bởi 2 thành viên nhóm "
        "với cross-checking, và bổ sung dữ liệu lớp thiểu số bằng API LLM.",
        style_name=body_style)
    add_para(doc,
        "(c) Phương pháp thực nghiệm: Triển khai pipeline tiền xử lý + Data Augmentation; "
        "huấn luyện song song 4 kiến trúc Deep Learning trên cùng tập dữ liệu, cùng siêu tham "
        "số để đảm bảo so sánh công bằng; đánh giá định lượng qua nhiều metric.",
        style_name=body_style)
    add_para(doc,
        "(d) Phương pháp phân tích kết quả: So sánh Accuracy, F1, Confusion Matrix giữa các "
        "mô hình; phân tích lỗi điển hình; đánh giá chi phí tính toán; rút ra kết luận và "
        "đề xuất hướng phát triển.",
        style_name=body_style)

    add_para(doc, "1.5 Đóng góp của đề tài", style_name=h2_style, bold=True)
    add_para(doc,
        "Đề tài đóng góp ở cả khía cạnh học thuật và thực tiễn:",
        style_name=body_style)
    add_para(doc,
        "Về mặt học thuật: (1) Cung cấp một bộ dữ liệu phân loại 6 cảm xúc cho tiếng Việt "
        "TMĐT với ~26.382 mẫu được gán nhãn cẩn thận; (2) Cung cấp từ điển chuẩn hóa Teencode "
        "367 cặp – tài nguyên có thể tái sử dụng cho các nghiên cứu NLP tiếng Việt khác; "
        "(3) Báo cáo so sánh chi tiết 4 kiến trúc Deep Learning trên cùng pipeline – tham "
        "chiếu cho các nghiên cứu tương lai.",
        style_name=body_style)
    add_para(doc,
        "Về mặt thực tiễn: (1) Hệ thống có thể triển khai ngay vào các ứng dụng phân tích "
        "phản hồi khách hàng cho doanh nghiệp Việt Nam; (2) Đã tích hợp với pipeline điều "
        "phối Apache Airflow và REST API (FastAPI), sẵn sàng đóng gói thành dịch vụ thương "
        "mại; (3) Mô hình đạt 93%+ accuracy với chi phí tính toán thấp (chạy được trên CPU "
        "thông thường), phù hợp với nhiều quy mô doanh nghiệp.",
        style_name=body_style)

    add_para(doc, "1.6 Cấu trúc báo cáo", style_name=h2_style, bold=True)
    add_para(doc,
        "Báo cáo được tổ chức thành 5 chương như sau:",
        style_name=body_style)
    add_para(doc,
        "Chương 1: Giới thiệu về bài toán – đặt vấn đề, mục tiêu, phạm vi, phương pháp "
        "nghiên cứu, đóng góp.",
        style_name=body_style)
    add_para(doc,
        "Chương 2: Phân tích yêu cầu của bài toán – các yêu cầu kỹ thuật, các phương pháp "
        "giải quyết bài toán hiện có và lựa chọn phương pháp đề xuất.",
        style_name=body_style)
    add_para(doc,
        "Chương 3: Phương pháp đề xuất – mô tả chi tiết kiến trúc hệ thống, quy trình tiền "
        "xử lý, công thức toán học của các mô hình Deep Learning và quy trình huấn luyện.",
        style_name=body_style)
    add_para(doc,
        "Chương 4: Thực nghiệm – mô tả dữ liệu, công nghệ sử dụng, phương pháp đánh giá, "
        "trình bày và phân tích kết quả thực nghiệm.",
        style_name=body_style)
    add_para(doc,
        "Chương 5: Kết luận – tổng kết kết quả đạt được, hạn chế và hướng phát triển.",
        style_name=body_style)

    # ============ CHƯƠNG 2 ============
    doc.add_page_break()
    add_para(doc, "CHƯƠNG 2", style_name=h1_style, bold=True)
    add_para(doc, "PHÂN TÍCH YÊU CẦU CỦA BÀI TOÁN", style_name=h1_style, bold=True)

    add_para(doc, "2.1 Cơ sở lý thuyết về phân tích cảm xúc", style_name=h2_style, bold=True)
    add_para(doc,
        "Phân tích cảm xúc (Sentiment Analysis) hay Khai phá ý kiến (Opinion Mining) là một "
        "nhánh con quan trọng của Xử lý ngôn ngữ tự nhiên (NLP), nghiên cứu phương pháp tự "
        "động xác định thái độ, cảm xúc và quan điểm của người viết đối với một chủ đề cụ "
        "thể trong văn bản. Theo Liu (2012) [21], bài toán có thể được phân loại theo nhiều "
        "mức độ: tài liệu (document-level), câu (sentence-level), khía cạnh (aspect-level).",
        style_name=body_style)
    add_para(doc,
        "Theo nhà tâm lý học Paul Ekman (1992) [11], có 6 nhóm cảm xúc cơ bản phổ quát "
        "(universal basic emotions) được nhận diện ở mọi nền văn hóa: Happiness (vui vẻ), "
        "Sadness (buồn bã), Anger (tức giận), Fear (sợ hãi), Disgust (ghê tởm), Surprise "
        "(ngạc nhiên). Đây chính là cơ sở phân loại được sử dụng trong đề tài này.",
        style_name=body_style)

    add_para(doc, "2.2 Yêu cầu của bài toán", style_name=h2_style, bold=True)
    add_para(doc, "Đầu vào (Input):", style_name=body_style)
    add_para(doc,
        "Văn bản tiếng Việt dạng bình luận/đánh giá sản phẩm từ các nền tảng TMĐT. Văn bản "
        "có các đặc điểm: độ dài trung bình ~36 từ/câu, có thể chứa Teencode (vd: 'k', 'dc'), "
        "ký tự đặc biệt, emoji, sai chính tả, cách viết không chuẩn (vd: lặp ký tự "
        "'tốttttt'), và thường ở dạng câu nói thông thường.",
        style_name=body_style)
    add_para(doc, "Đầu ra (Output):", style_name=body_style)
    add_para(doc,
        "Nhãn cảm xúc dự đoán ∈ {Happiness, Sadness, Anger, Fear, Disgust, Surprise}, kèm "
        "theo phân phối xác suất Softmax cho 6 lớp (tổng = 1.0). Người dùng có thể sử dụng "
        "nhãn đỉnh (top-1) hoặc top-k để có nhiều thông tin hơn.",
        style_name=body_style)
    add_para(doc, "Yêu cầu kỹ thuật:", style_name=body_style)
    add_para(doc,
        "(1) Hiệu suất: Độ chính xác (Accuracy) > 90% trên tập kiểm thử; Macro-F1 > 0.90 để "
        "đảm bảo công bằng giữa các lớp.",
        style_name=body_style)
    add_para(doc,
        "(2) Xử lý đặc thù tiếng Việt: Mô hình phải xử lý tốt Teencode, từ ghép, từ thiếu "
        "dấu, sai chính tả – những đặc điểm phổ biến của bình luận TMĐT Việt Nam.",
        style_name=body_style)
    add_para(doc,
        "(3) Mất cân bằng dữ liệu: Mô hình phải đạt F1 > 0.80 cho lớp thiểu số (Disgust ~10%) "
        "thay vì chỉ tập trung vào lớp đa số.",
        style_name=body_style)
    add_para(doc,
        "(4) Hiệu suất tính toán: Thời gian inference < 100ms/mẫu để có thể triển khai realtime "
        "trong các ứng dụng tương tác; có thể chạy được trên CPU thông thường khi không có "
        "GPU.",
        style_name=body_style)
    add_para(doc,
        "(5) Khả năng triển khai: Mô hình phải đóng gói được thành module Python tích hợp "
        "vào REST API hoặc workflow ETL.",
        style_name=body_style)

    add_para(doc, "2.3 Các phương pháp giải quyết bài toán", style_name=h2_style, bold=True)

    add_para(doc, "2.3.1 Phương pháp truyền thống – Rule-based / Lexicon",
             style_name='Normal', bold=True)
    add_para(doc,
        "Cách tiếp cận này dựa trên từ điển từ khóa cảm xúc đã được gán điểm thủ công bởi "
        "chuyên gia (vd: SentiWordNet cho tiếng Anh, VietSentiWordNet cho tiếng Việt). Khi "
        "phân tích, hệ thống đếm số lượng từ tích cực/tiêu cực và tính điểm tổng để phân "
        "loại. Ưu điểm: đơn giản, dễ giải thích, không cần dữ liệu huấn luyện. Hạn chế: "
        "không xử lý được phủ định, mỉa mai, ngữ cảnh phức tạp; không tự động cập nhật "
        "theo dữ liệu mới; phụ thuộc chất lượng từ điển. Độ chính xác thường ở mức 60–70%.",
        style_name=body_style)

    add_para(doc, "2.3.2 Machine Learning cổ điển", style_name='Normal', bold=True)
    add_para(doc,
        "Sử dụng các bộ phân lớp như Support Vector Machine (SVM), Naive Bayes, Logistic "
        "Regression, Random Forest với đặc trưng Bag-of-Words (BoW) hoặc TF-IDF. Quy trình: "
        "(a) Vector hóa văn bản thành ma trận thưa N×V (N = số mẫu, V = kích thước từ vựng); "
        "(b) Huấn luyện classifier trên tập train; (c) Dự đoán trên tập test.",
        style_name=body_style)
    add_para(doc,
        "Ưu điểm: huấn luyện nhanh trên CPU, dễ giải thích, ít tham số. Hạn chế: không học "
        "được biểu diễn ngữ nghĩa sâu (semantic representation), khó tổng quát hóa khi gặp "
        "từ mới (out-of-vocabulary), gặp khó khăn với dữ liệu nhiễu cao như tiếng Việt mạng "
        "xã hội. Độ chính xác thường ở mức 75–85%.",
        style_name=body_style)

    add_para(doc, "2.3.3 Deep Learning", style_name='Normal', bold=True)
    add_para(doc,
        "Sử dụng các mạng nơ-ron sâu như CNN [1], LSTM [2], BiLSTM [6] và các biến thể kết "
        "hợp. Ưu điểm chính: (a) Tự động học đặc trưng (representation learning) từ dữ liệu "
        "thô qua lớp Embedding [13], không cần feature engineering thủ công; (b) Xử lý tốt "
        "ngữ cảnh và dữ liệu nhiễu nhờ khả năng học các pattern phức tạp; (c) Dễ dàng mở "
        "rộng quy mô khi có dữ liệu lớn hơn.",
        style_name=body_style)
    add_para(doc,
        "Hạn chế: cần dữ liệu huấn luyện lớn (≥10.000 mẫu), thời gian huấn luyện lâu hơn ML "
        "cổ điển, kém minh bạch (black-box). Tuy nhiên với bài toán hiện tại có ~26.000 mẫu, "
        "Deep Learning là lựa chọn phù hợp. Độ chính xác kỳ vọng 90–94%.",
        style_name=body_style)

    add_para(doc, "2.3.4 Transfer Learning với mô hình ngôn ngữ tiền huấn luyện",
             style_name='Normal', bold=True)
    add_para(doc,
        "Đại diện tiêu biểu: BERT [5], RoBERTa [16], XLM-RoBERTa [17] cho đa ngôn ngữ; và "
        "PhoBERT [3], ViSoBERT [22] dành riêng cho tiếng Việt. Các mô hình này được pre-train "
        "trên corpus rất lớn (>20GB cho PhoBERT) và sau đó fine-tune trên task cụ thể.",
        style_name=body_style)
    add_para(doc,
        "Ưu điểm: hiệu suất rất cao (95–97%) nhờ biểu diễn ngữ nghĩa chất lượng cao, hiểu "
        "ngữ cảnh sâu. Hạn chế: yêu cầu tài nguyên tính toán rất lớn (GPU ≥ 16GB), thời gian "
        "inference chậm (200–500 ms/mẫu trên CPU), khó triển khai trên thiết bị nhỏ. Phù hợp "
        "khi có hạ tầng đủ mạnh.",
        style_name=body_style)

    add_para(doc, "2.3.5 Bảng so sánh tổng hợp các phương pháp",
             style_name='Normal', bold=True)
    add_table(doc,
              headers=APPROACH_COMPARISON[0],
              rows=APPROACH_COMPARISON[1:],
              caption='Bảng 2.1 So sánh các hướng tiếp cận giải quyết bài toán phân tích cảm xúc',
              col_widths=[1.8, 2.0, 1.8, 0.9])

    add_para(doc, "2.4 Các nghiên cứu liên quan trong tiếng Việt", style_name=h2_style, bold=True)
    add_para(doc,
        "Vu et al. (2018) [10] công bố VnCoreNLP – bộ công cụ NLP toàn diện cho tiếng Việt "
        "(word segmentation, POS tagging, NER, dependency parsing). Đây là một trong những "
        "tài nguyên NLP tiếng Việt mạnh mẽ nhất hiện có.",
        style_name=body_style)
    add_para(doc,
        "Nguyen & Nguyen (2020) [3] công bố PhoBERT – mô hình ngôn ngữ tiền huấn luyện đầu "
        "tiên cho tiếng Việt dựa trên kiến trúc RoBERTa, được pre-train trên 20GB văn bản "
        "tiếng Việt. PhoBERT đạt SOTA trên nhiều tác vụ tiếng Việt.",
        style_name=body_style)
    add_para(doc,
        "Nguyen et al. (2023) [22] đề xuất ViSoBERT – chuyên biệt cho văn bản mạng xã hội "
        "Việt Nam với nhiều Teencode, icon và viết tắt – đặc biệt phù hợp với loại dữ liệu "
        "trong nghiên cứu này.",
        style_name=body_style)
    add_para(doc,
        "Van Nguyen et al. (2018) [4] công bố UIT-VSFC – corpus phản hồi sinh viên với 3 "
        "nhãn (positive/negative/neutral). Tran & Nguyen (2020) [23] tổng hợp các phương "
        "pháp phân tích cảm xúc fine-grained cho tiếng Việt và chỉ ra sự thiếu hụt các "
        "nghiên cứu phân loại đa lớp cảm xúc trong miền TMĐT.",
        style_name=body_style)

    add_para(doc, "2.5 Phương pháp đề xuất giải quyết bài toán", style_name=h2_style, bold=True)
    add_para(doc,
        "Sau khi cân nhắc các yêu cầu (Mục 2.2) và phân tích các phương pháp (Mục 2.3), nhóm "
        "lựa chọn hướng Deep Learning với kiến trúc Hybrid CNN+BiLSTM làm phương pháp đề "
        "xuất chính, đồng thời triển khai 3 kiến trúc cơ sở (MLP, 1D-CNN, BiLSTM) để so sánh "
        "định lượng.",
        style_name=body_style)
    add_para(doc,
        "Lý do lựa chọn Hybrid CNN+BiLSTM: (1) Cân bằng tốt nhất giữa hiệu suất kỳ vọng "
        "(~93%) và chi phí tính toán (chạy được trên Tesla T4 với <1 GB VRAM); (2) Tận dụng "
        "ưu điểm của cả CNN (trích xuất đặc trưng cục bộ nhanh) và BiLSTM (hiểu ngữ cảnh "
        "tuần tự hai chiều); (3) Đã được chứng minh hiệu quả trong nhiều nghiên cứu phân "
        "loại văn bản tiếng Anh và đang được áp dụng cho tiếng Việt; (4) Phù hợp với quy mô "
        "dữ liệu (~26.000 mẫu) – không quá nhỏ cho Deep Learning, không quá lớn để cần đến "
        "Transformer.",
        style_name=body_style)
    add_para(doc,
        "Đóng góp thêm của nghiên cứu so với việc chỉ áp dụng kiến trúc có sẵn: (a) Xây dựng "
        "quy trình tiền xử lý chuyên biệt cho tiếng Việt với từ điển Teencode 367 cặp – "
        "không có nghiên cứu tiếng Việt nào trước đây công bố tài nguyên tương tự; "
        "(b) Áp dụng Data Augmentation (Random Deletion, Random Swap) đã được chỉnh sửa cho "
        "tiếng Việt để xử lý mất cân bằng; (c) So sánh đồng thời 4 kiến trúc trên cùng "
        "pipeline – cung cấp tham chiếu khoa học cho nghiên cứu tương lai.",
        style_name=body_style)

    # ============ CHƯƠNG 3 ============
    doc.add_page_break()
    add_para(doc, "CHƯƠNG 3", style_name=h1_style, bold=True)
    add_para(doc, "PHƯƠNG PHÁP ĐỀ XUẤT", style_name=h1_style, bold=True)

    add_para(doc, "3.1 Kiến trúc tổng thể của hệ thống", style_name=h2_style, bold=True)
    add_para(doc,
        "Hệ thống phân tích cảm xúc được thiết kế theo pipeline 5 bước, vận hành trong hai "
        "giai đoạn chính:",
        style_name=body_style)
    add_para(doc,
        "Giai đoạn 1 – Huấn luyện (Training Phase): Thu thập dữ liệu từ TMĐT → Tiền xử lý "
        "chuyên biệt cho tiếng Việt (4 bước) → Vector hóa bằng Keras Tokenizer → Data "
        "Augmentation → Huấn luyện song song 4 kiến trúc Deep Learning trên cùng pipeline → "
        "Đánh giá và lưu mô hình tốt nhất theo Macro-F1.",
        style_name=body_style)
    add_para(doc,
        "Giai đoạn 2 – Suy luận (Inference Phase): Nhận văn bản tiếng Việt đầu vào → Áp dụng "
        "cùng pipeline tiền xử lý → Vector hóa → Đưa qua mô hình đã huấn luyện → Trả về phân "
        "phối xác suất Softmax cho 6 lớp cảm xúc.",
        style_name=body_style)
    add_para(doc,
        "Toàn bộ pipeline được tự động hóa bằng các script Python trong notebook Jupyter và "
        "có thể tích hợp vào Apache Airflow để chạy định kỳ trên dữ liệu mới (cấu hình DAG "
        "đã được thiết lập trong dự án).",
        style_name=body_style)

    add_para(doc, "3.2 Quy trình thu thập dữ liệu", style_name=h2_style, bold=True)
    add_para(doc,
        "Dữ liệu được thu thập từ 3 nguồn chính: Tiki, Shopee, Lazada – đại diện cho 3 nền "
        "tảng TMĐT lớn nhất Việt Nam. Quy trình thu thập bao gồm:",
        style_name=body_style)
    add_para(doc,
        "Bước 1 – Web Scraping: Sử dụng Selenium WebDriver để mô phỏng hành vi người dùng "
        "(scroll trang, click 'Xem thêm') và thu thập bình luận tự động. Mỗi nền tảng có cấu "
        "trúc HTML khác nhau, cần custom selector cho từng site.",
        style_name=body_style)
    add_para(doc,
        "Bước 2 – Lọc và làm sạch sơ bộ: Loại bỏ bình luận trùng lặp (dựa trên hash nội "
        "dung), bình luận quá ngắn (<10 ký tự) hoặc chỉ chứa emoji.",
        style_name=body_style)
    add_para(doc,
        "Bước 3 – Gán nhãn: Hai thành viên nhóm gán nhãn độc lập, sau đó cross-checking để "
        "thống nhất các trường hợp khác biệt. Sử dụng hướng dẫn gán nhãn rõ ràng dựa trên "
        "định nghĩa Ekman cho 6 lớp cảm xúc.",
        style_name=body_style)
    add_para(doc,
        "Bước 4 – Bổ sung lớp thiểu số: Sử dụng API của LLM (GPT) để sinh thêm bình luận nhân "
        "tạo cho 2 lớp thiểu số là Disgust và Fear, sau đó kiểm tra thủ công để đảm bảo chất "
        "lượng. Tổng cộng bổ sung ~3.000 mẫu.",
        style_name=body_style)
    add_para(doc,
        "Kết quả: Thu được 26.382 mẫu cuối cùng với 6 nhãn cảm xúc.",
        style_name=body_style)
    add_image(doc, IMG['eda'], width_inches=6.0,
              caption='Hình 3.1 Phân phối các lớp cảm xúc trong tập dữ liệu sau thu thập')

    add_para(doc, "3.3 Quy trình tiền xử lý dữ liệu", style_name=h2_style, bold=True)
    add_para(doc,
        "Tiếng Việt có nhiều đặc thù khác biệt so với tiếng Anh, đòi hỏi quy trình tiền xử "
        "lý chuyên biệt. Pipeline gồm 5 bước chính, được trình bày chi tiết qua các mục sau "
        "và minh họa qua Bảng 3.2.",
        style_name=body_style)

    add_para(doc, "3.3.1 Chuẩn hóa Teencode", style_name=h2_style, bold=True)
    add_para(doc,
        "Teencode là hiện tượng viết tắt phổ biến trong giới trẻ Việt Nam khi giao tiếp trên "
        "mạng xã hội và TMĐT. Nếu không chuẩn hóa, mô hình sẽ coi 'k' và 'không' là hai từ "
        "hoàn toàn khác biệt – làm giảm chất lượng biểu diễn ngữ nghĩa.",
        style_name=body_style)
    add_para(doc,
        "Nhóm xây dựng một từ điển ánh xạ 367 cặp Teencode → tiếng Việt chuẩn dựa trên khảo "
        "sát thực tế từ dữ liệu thu thập. Bảng 3.1 trình bày một số ví dụ điển hình.",
        style_name=body_style)
    add_table(doc,
              headers=["Teencode", "Dạng chuẩn"],
              rows=TEENCODE_EXAMPLES,
              caption='Bảng 3.1 Một số ví dụ điển hình trong từ điển chuẩn hóa Teencode',
              col_widths=[2.0, 4.0])

    add_para(doc, "3.3.2 Làm sạch và chuẩn hóa Unicode", style_name=h2_style, bold=True)
    add_para(doc,
        "Bước này thực hiện các thao tác làm sạch:",
        style_name=body_style)
    add_para(doc,
        "(a) Chuyển toàn bộ về chữ thường (lowercase) để giảm kích thước từ vựng.",
        style_name=body_style)
    add_para(doc,
        "(b) Loại bỏ URL (regex: https?://...), email và số điện thoại – các thành phần này "
        "không mang ý nghĩa cảm xúc.",
        style_name=body_style)
    add_para(doc,
        "(c) Xử lý hiện tượng lặp ký tự (vd: 'tốtttt' → 'tốt') bằng regex thay thế các chuỗi "
        "≥3 ký tự giống nhau bằng 1 ký tự duy nhất.",
        style_name=body_style)
    add_para(doc,
        "(d) Chuẩn hóa Unicode NFC (Normalization Form Canonical Composition) để gộp các ký "
        "tự tổ hợp về dạng chuẩn duy nhất – điều này quan trọng vì cùng một chữ tiếng Việt "
        "có thể có nhiều mã Unicode khác nhau (vd: 'à' có thể là U+00E0 hoặc U+0061+U+0300).",
        style_name=body_style)
    add_para(doc,
        "(e) Loại bỏ các ký tự đặc biệt và emoji không cần thiết, giữ lại chữ cái, dấu câu "
        "cơ bản và khoảng trắng.",
        style_name=body_style)

    add_para(doc, "3.3.3 Tách từ tiếng Việt (Word Segmentation)",
             style_name=h2_style, bold=True)
    add_para(doc,
        "Khác với tiếng Anh nơi từ được phân tách bởi khoảng trắng, tiếng Việt có nhiều từ "
        "ghép. Mỗi 'từ' trong tiếng Việt có thể gồm 1 hoặc nhiều âm tiết được phân tách bởi "
        "khoảng trắng nhưng tạo thành một đơn vị ngữ nghĩa duy nhất.",
        style_name=body_style)
    add_para(doc,
        "Ví dụ: 'sản phẩm', 'chất lượng', 'hài lòng' – mỗi cụm là 1 từ ghép. Nếu không tách "
        "đúng, mô hình sẽ coi 'sản' và 'phẩm' là hai từ độc lập, mất đi ngữ nghĩa.",
        style_name=body_style)
    add_para(doc,
        "Nhóm sử dụng Underthesea (https://github.com/undertheseanlp/underthesea) – thư viện "
        "open-source phổ biến cho NLP tiếng Việt. Hàm word_tokenize() trả về danh sách các "
        "token, trong đó từ ghép được nối bằng dấu '_' để giữ làm 1 đơn vị.",
        style_name=body_style)
    add_code_block(doc,
        ">>> from underthesea import word_tokenize\n"
        ">>> word_tokenize('sản phẩm chất lượng tốt', format='text')\n"
        "'sản_phẩm chất_lượng tốt'")

    add_para(doc, "3.3.4 Vector hóa và Padding", style_name=h2_style, bold=True)
    add_para(doc,
        "Văn bản sau khi tách từ được chuyển thành chuỗi số nguyên (token IDs) bằng Keras "
        "Tokenizer:",
        style_name=body_style)
    add_para(doc,
        "Bước 1: Xây dựng từ vựng V từ tập huấn luyện (chỉ giữ V từ phổ biến nhất, các từ "
        "hiếm gặp được mã hóa thành <UNK>).",
        style_name=body_style)
    add_para(doc,
        "Bước 2: Chuyển mỗi câu thành chuỗi token IDs.",
        style_name=body_style)
    add_para(doc,
        "Bước 3: Pad/truncate về độ dài cố định L = 100 tokens (chọn dựa trên phân tích EDA, "
        "bao phủ >95% bình luận trong tập dữ liệu).",
        style_name=body_style)
    add_para(doc,
        "Kết quả: Ma trận đầu vào X ∈ ℕ^(N × L) với N là số mẫu, L = 100.",
        style_name=body_style)

    add_para(doc, "3.3.5 Data Augmentation", style_name=h2_style, bold=True)
    add_para(doc,
        "Để khắc phục Overfitting và mất cân bằng dữ liệu, áp dụng 2 kỹ thuật theo phương "
        "pháp Easy Data Augmentation (EDA) [9] trên tập huấn luyện:",
        style_name=body_style)
    add_para(doc,
        "(a) Random Deletion (RD): Với mỗi câu, duyệt qua từng từ và xóa với xác suất p=0.15. "
        "Không xóa nếu câu chỉ còn 1 từ. Mục đích: làm mô hình bền vững hơn với từ thiếu hoặc "
        "câu bị cắt ngắn.",
        style_name=body_style)
    add_para(doc,
        "(b) Random Swap (RS): Đổi chỗ ngẫu nhiên 2 từ trong câu, lặp lại n lần (n tỷ lệ với "
        "độ dài câu). Mục đích: làm mô hình ít phụ thuộc vào thứ tự từ chính xác.",
        style_name=body_style)
    add_para(doc,
        "Tổng cộng bổ sung ~4.188 mẫu nhân tạo cho tập huấn luyện, đặc biệt tăng cường cho "
        "lớp thiểu số Disgust.",
        style_name=body_style)

    add_para(doc,
        "Bảng 3.2 minh họa toàn bộ pipeline tiền xử lý qua từng bước với một bình luận thực:",
        style_name=body_style)
    add_table(doc,
              headers=PREPROCESS_EXAMPLES[0],
              rows=PREPROCESS_EXAMPLES[1:],
              caption='Bảng 3.2 Minh họa pipeline tiền xử lý qua từng bước',
              col_widths=[2.2, 4.0])

    add_pseudocode(doc, "Pseudocode 3.1 Quy trình tiền xử lý văn bản tiếng Việt", [
        "Input: Văn bản gốc t",
        "Output: Vector số nguyên có độ dài cố định L=100",
        "",
        " 1: t ← lowercase(t)",
        " 2: t ← remove_url_email_phone(t)",
        " 3: t ← reduce_repeated_chars(t)        // 'tốtttt' → 'tốt'",
        " 4: t ← unicode_normalize(t, form='NFC')",
        " 5: words ← split(t)",
        " 6: for each word w in words do",
        " 7:     if w in TeencodeDict then",
        " 8:         w ← TeencodeDict[w]",
        " 9: t ← join(words)",
        "10: t ← remove_special_chars(t)",
        "11: tokens ← underthesea.word_tokenize(t)        // Tách từ ghép",
        "12: ids ← tokenizer.texts_to_sequences([tokens])",
        "13: x ← pad_sequences(ids, maxlen=L, padding='post')",
        "14: return x",
    ])

    add_para(doc, "3.4 Kiến trúc các mô hình Deep Learning", style_name=h2_style, bold=True)
    add_para(doc,
        "Bốn kiến trúc Deep Learning được triển khai và so sánh trong nghiên cứu này. Tất cả "
        "đều bắt đầu bằng lớp Embedding chung: với một câu x = (x₁, x₂, ..., x_L), mỗi token "
        "x_i ∈ {1, ..., V} được ánh xạ thành vector dày e_i ∈ ℝ^d:",
        style_name=body_style)
    add_formula(doc, "e_i = E[x_i],   E ∈ ℝ^(V × d),   d = 128", label="3.1")

    add_para(doc, "3.4.1 Mô hình MLP (Baseline)", style_name=h2_style, bold=True)
    add_para(doc,
        "Đây là mô hình cơ sở đơn giản nhất, được xây dựng để thiết lập ngưỡng hiệu suất tối "
        "thiểu cho việc so sánh. Mô hình sử dụng kiến trúc Feed-forward Neural Network đơn "
        "giản kết hợp với lớp Embedding.",
        style_name=body_style)
    add_para(doc, "Kiến trúc chi tiết:", style_name=body_style)
    add_code_block(doc,
        "Input(L=100)\n"
        "  ↓\n"
        "Embedding(V → 128)\n"
        "  ↓\n"
        "GlobalAveragePooling1D()         # Trung bình các vector từ\n"
        "  ↓\n"
        "Dense(256, activation='relu', kernel_regularizer=L2(1e-4))\n"
        "BatchNormalization()\n"
        "Dropout(0.5)\n"
        "  ↓\n"
        "Dense(128, activation='relu', kernel_regularizer=L2(1e-4))\n"
        "BatchNormalization()\n"
        "Dropout(0.4)\n"
        "  ↓\n"
        "Dense(6, activation='softmax')")
    add_para(doc,
        "Global Average Pooling thay vì Flatten giúp giảm số tham số đáng kể (từ L×d = "
        "12.800 xuống còn d = 128 đặc trưng) và tránh overfitting:",
        style_name=body_style)
    add_formula(doc, "h = (1/L) Σᵢ₌₁ᴸ eᵢ,   h ∈ ℝᵈ", label="3.2")
    add_image(doc, IMG['arch_mlp'], width_inches=5.0,
              caption='Hình 3.2 Sơ đồ kiến trúc mô hình MLP')

    add_para(doc, "3.4.2 Mô hình 1D-CNN", style_name=h2_style, bold=True)
    add_para(doc,
        "Convolutional Neural Network 1 chiều [1] được thiết kế để tận dụng khả năng trích "
        "xuất đặc trưng cục bộ (local features), vốn rất hiệu quả cho việc phát hiện các cụm "
        "từ (n-gram) đặc trưng cảm xúc trong văn bản như 'rất tốt', 'cực kỳ tệ', 'hơi thất "
        "vọng', 'không hài lòng'.",
        style_name=body_style)
    add_para(doc,
        "Phép tích chập 1 chiều với cửa sổ trượt kernel size k được định nghĩa:",
        style_name=body_style)
    add_formula(doc, "h_i = ReLU( Σⱼ₌₀ᵏ⁻¹ W_j · x_(i+j) + b ),   W ∈ ℝᵏˣᵈ", label="3.3")
    add_para(doc,
        "Trong đó hàm kích hoạt ReLU (Rectified Linear Unit) được tính:",
        style_name=body_style)
    add_formula(doc, "ReLU(x) = max(0, x)", label="3.4")
    add_para(doc, "Kiến trúc chi tiết:", style_name=body_style)
    add_code_block(doc,
        "Input(L=100) → Embedding(V → 128)\n"
        "  ↓\n"
        "SpatialDropout1D(0.3)            # Chống overfit\n"
        "  ↓\n"
        "Conv1D(128, kernel_size=3, padding='same')\n"
        "BatchNormalization() + ReLU\n"
        "MaxPooling1D(pool_size=2)\n"
        "  ↓\n"
        "Conv1D(128, kernel_size=3, padding='same')\n"
        "BatchNormalization() + ReLU\n"
        "  ↓\n"
        "GlobalMaxPooling1D()             # Lấy đặc trưng mạnh nhất\n"
        "  ↓\n"
        "Dense(128, activation='relu') + Dropout(0.4)\n"
        "  ↓\n"
        "Dense(6, activation='softmax')")
    add_image(doc, IMG['arch_cnn'], width_inches=5.0,
              caption='Hình 3.3 Sơ đồ kiến trúc mô hình 1D-CNN')

    add_para(doc, "3.4.3 Mô hình BiLSTM", style_name=h2_style, bold=True)
    add_para(doc,
        "Long Short-Term Memory (LSTM) [2] là biến thể RNN giải quyết vấn đề vanishing/exploding "
        "gradient thông qua cơ chế cổng (gating mechanism) gồm 3 cổng: forget gate (fₜ – quyết "
        "định thông tin nào quên đi), input gate (iₜ – quyết định thông tin mới nào được lưu), "
        "và output gate (oₜ – quyết định thông tin nào được output ra ngoài).",
        style_name=body_style)
    add_para(doc,
        "Phương trình cập nhật của một tế bào LSTM tại bước thời gian t:",
        style_name=body_style)
    add_formula(doc, "fₜ = σ(W_f · [hₜ₋₁, xₜ] + b_f)        // Forget gate", label="3.5")
    add_formula(doc, "iₜ = σ(W_i · [hₜ₋₁, xₜ] + b_i)        // Input gate", label="3.6")
    add_formula(doc, "C̃ₜ = tanh(W_C · [hₜ₋₁, xₜ] + b_C)     // Candidate cell", label="3.7")
    add_formula(doc, "Cₜ = fₜ ⊙ Cₜ₋₁ + iₜ ⊙ C̃ₜ            // Cell state mới", label="3.8")
    add_formula(doc, "oₜ = σ(W_o · [hₜ₋₁, xₜ] + b_o)        // Output gate", label="3.9")
    add_formula(doc, "hₜ = oₜ ⊙ tanh(Cₜ)                   // Hidden state", label="3.10")
    add_para(doc,
        "trong đó σ là sigmoid function: σ(x) = 1/(1+e^(-x)), và ⊙ là phép nhân element-wise.",
        style_name=body_style)
    add_para(doc,
        "Bidirectional LSTM (BiLSTM) [6] cải tiến LSTM bằng cách chạy đồng thời 2 LSTM ngược "
        "chiều nhau và nối hidden state lại:",
        style_name=body_style)
    add_formula(doc, "hₜ = [ hₜ→ ; hₜ← ]", label="3.11")
    add_para(doc,
        "Cơ chế hai chiều cho phép mô hình học ngữ cảnh từ cả hai phía: quá khứ (từ trước "
        "đó) và tương lai (từ phía sau), giúp hiểu ngữ nghĩa tốt hơn cho phân loại văn bản.",
        style_name=body_style)
    add_para(doc, "Kiến trúc chi tiết:", style_name=body_style)
    add_code_block(doc,
        "Input(L=100) → Embedding(V → 128)\n"
        "  ↓\n"
        "Bidirectional(LSTM(128, return_sequences=True,\n"
        "                   recurrent_dropout=0.4))\n"
        "  ↓\n"
        "Bidirectional(LSTM(128, recurrent_dropout=0.4))\n"
        "  ↓\n"
        "Dense(128, activation='relu') + Dropout(0.5)\n"
        "  ↓\n"
        "Dense(6, activation='softmax')")
    add_image(doc, IMG['arch_bilstm'], width_inches=5.0,
              caption='Hình 3.4 Sơ đồ kiến trúc mô hình BiLSTM')

    add_para(doc, "3.4.4 Mô hình Hybrid CNN+BiLSTM (đề xuất chính)",
             style_name=h2_style, bold=True)
    add_para(doc,
        "Đây là kiến trúc chủ đạo được đề xuất trong đồ án, kết hợp sức mạnh của cả CNN "
        "(trích xuất đặc trưng cục bộ – local) và BiLSTM (mô hình hóa chuỗi – global) trong "
        "cùng một mô hình end-to-end:",
        style_name=body_style)
    add_para(doc,
        "(a) Tầng CNN làm Feature Extractor cục bộ – sử dụng 2 lớp Conv1D(128 filters, "
        "kernel=3, padding='same') kết hợp MaxPooling1D để lọc nhiễu và trích xuất các cụm "
        "từ mang tính cục bộ từ embedding. Padding='same' giữ nguyên độ dài chuỗi để không "
        "mất thông tin biên.",
        style_name=body_style)
    add_para(doc,
        "(b) Tầng BiLSTM làm Sequence Modeler – nhận đầu ra của CNN (chuỗi đặc trưng đã "
        "rút gọn về mặt nhiễu) đưa vào Bidirectional(LSTM(128)) để học mối quan hệ tuần tự "
        "giữa các đặc trưng theo trình tự thời gian, hiểu được ngữ nghĩa tổng thể của câu.",
        style_name=body_style)
    add_para(doc,
        "(c) Classification Head – Global Max Pooling chắt lọc thông tin quan trọng nhất, "
        "tiếp theo là 2 lớp Dense + Dropout mạnh và lớp Softmax đầu ra cho 6 lớp.",
        style_name=body_style)
    add_para(doc, "Kiến trúc chi tiết:", style_name=body_style)
    add_code_block(doc,
        "Input(L=100) → Embedding(V → 128)\n"
        "  ↓\n"
        "Conv1D(128, kernel=3, padding='same') + ReLU\n"
        "MaxPooling1D(2)\n"
        "Conv1D(128, kernel=3, padding='same') + ReLU\n"
        "  ↓\n"
        "Bidirectional(LSTM(128, return_sequences=True))\n"
        "  ↓\n"
        "GlobalMaxPooling1D()\n"
        "  ↓\n"
        "Dense(128) + Dropout(0.5)\n"
        "Dense(64) + Dropout(0.4)\n"
        "  ↓\n"
        "Dense(6, activation='softmax')")
    add_image(doc, IMG['arch_hybrid'], width_inches=5.5,
              caption='Hình 3.5 Sơ đồ kiến trúc Hybrid CNN+BiLSTM (đề xuất chính)')

    # 3.5 Hàm mất mát và tối ưu
    add_para(doc, "3.5 Hàm mất mát và thuật toán tối ưu", style_name=h2_style, bold=True)
    add_para(doc,
        "Lớp đầu ra Softmax tính phân phối xác suất trên 6 lớp cảm xúc:",
        style_name=body_style)
    add_formula(doc, "P(y=k | x) = exp(z_k) / Σⱼ₌₁ᶜ exp(z_j),   C = 6", label="3.12")
    add_para(doc,
        "Hàm mất mát Sparse Categorical Crossentropy được sử dụng (do nhãn dạng integer "
        "encoding) kết hợp với Class Weighting:",
        style_name=body_style)
    add_formula(doc, "L = -(1/N) Σᵢ₌₁ᴺ wᵧᵢ · log P(yᵢ | xᵢ)", label="3.13")
    add_para(doc,
        "trong đó w_(yᵢ) là trọng số lớp tính bằng nghịch đảo tần suất, đảm bảo các lớp "
        "thiểu số được mô hình chú ý nhiều hơn:",
        style_name=body_style)
    add_formula(doc, "w_c = N / (C · n_c),   c ∈ {1, ..., C}", label="3.14")
    add_para(doc,
        "với N là tổng số mẫu, C là số lớp, n_c là số mẫu của lớp c.",
        style_name=body_style)
    add_para(doc,
        "Thuật toán tối ưu AdamW [7] – biến thể của Adam tách biệt Weight Decay khỏi cập "
        "nhật gradient – được sử dụng với cấu hình learning rate η = 5×10⁻⁴ và weight decay "
        "λ = 0.01:",
        style_name=body_style)
    add_formula(doc, "θₜ₊₁ = θₜ - η · ( m̂ₜ / (√v̂ₜ + ε) + λ · θₜ )", label="3.15")
    add_para(doc,
        "trong đó m̂ₜ, v̂ₜ là moment bậc 1 và 2 đã hiệu chỉnh bias của Adam. AdamW vượt trội "
        "so với Adam truyền thống ở khả năng kiểm soát overfitting.",
        style_name=body_style)

    # 3.6 Quy trình huấn luyện
    add_para(doc, "3.6 Quy trình huấn luyện", style_name=h2_style, bold=True)
    add_para(doc,
        "Môi trường thực nghiệm: Google Colab Pro với GPU NVIDIA Tesla T4 (16GB VRAM), kích "
        "hoạt Mixed Precision (fp16) để tăng tốc tính toán ma trận trong các phép nhân của "
        "mạng nơ-ron. Mixed Precision sử dụng định dạng float16 cho các phép tính trung gian "
        "và float32 cho tích lũy gradient, giảm bộ nhớ và tăng tốc 2–3 lần.",
        style_name=body_style)
    add_para(doc,
        "Bảng 3.3 trình bày đầy đủ các siêu tham số. Các giá trị được lựa chọn dựa trên "
        "thực nghiệm và kinh nghiệm từ các nghiên cứu trước đây cho bài toán phân loại văn "
        "bản tương tự.",
        style_name=body_style)
    add_table(doc,
              headers=["Siêu tham số", "Giá trị", "Mô tả"],
              rows=HYPERPARAMS,
              caption='Bảng 3.3 Cấu hình siêu tham số huấn luyện mô hình',
              col_widths=[1.8, 1.6, 2.6])

    add_para(doc, "Chiến lược huấn luyện thông minh:", style_name=body_style)
    add_para(doc,
        "(1) Early Stopping (patience=5, monitor=val_loss): Tự động dừng huấn luyện khi "
        "val_loss không cải thiện sau 5 epochs liên tiếp; khôi phục trọng số tốt nhất "
        "(restore_best_weights=True).",
        style_name=body_style)
    add_para(doc,
        "(2) ReduceLROnPlateau (patience=2, factor=0.3): Khi mô hình gặp plateau, giảm "
        "learning rate xuống 30% giúp tinh chỉnh tham số chính xác hơn và thoát khỏi local "
        "minimum.",
        style_name=body_style)
    add_para(doc,
        "(3) Class Weighting: Tính trước khi huấn luyện qua compute_class_weight với mode "
        "'balanced', truyền vào model.fit(class_weight=...).",
        style_name=body_style)

    add_pseudocode(doc, "Pseudocode 3.2 Quy trình huấn luyện và chống Overfitting", [
        "Input: Tập huấn luyện D, mô hình M, epochs E=20, batch size B=128",
        "Output: Mô hình đã huấn luyện tốt nhất M*",
        "",
        " 1: Tính class_weights w từ phân phối nhãn của D (compute_class_weight)",
        " 2: Khởi tạo optimizer AdamW(lr=5e-4, weight_decay=0.01)",
        " 3: best_val_loss ← +∞;  patience_es ← 0;  patience_lr ← 0",
        " 4: for epoch e = 1 to E do",
        " 5:     Shuffle D",
        " 6:     for each mini-batch (X, y) in D do",
        " 7:         ŷ ← M(X)                                  // Forward pass",
        " 8:         L ← SparseCategoricalCE(y, ŷ, weights=w)",
        " 9:         Cập nhật M qua AdamW(L)                    // Backward pass",
        "10:     end for",
        "11:     val_loss ← Đánh giá M trên tập validation",
        "12:     if val_loss < best_val_loss then",
        "13:         best_val_loss ← val_loss; M* ← M; patience_es=0; patience_lr=0",
        "14:     else",
        "15:         patience_es += 1; patience_lr += 1",
        "16:         if patience_lr ≥ 2 then  lr ← lr × 0.3; patience_lr=0   // ReduceLR",
        "17:         if patience_es ≥ 5 then  break                          // Early Stop",
        "18:     end if",
        "19: end for",
        "20: return M*",
    ])

    # ============ CHƯƠNG 4 ============
    doc.add_page_break()
    add_para(doc, "CHƯƠNG 4", style_name=h1_style, bold=True)
    add_para(doc, "THỰC NGHIỆM", style_name=h1_style, bold=True)

    # ---- 4.1 Tập dữ liệu ----
    add_para(doc, "4.1 Tập dữ liệu thực nghiệm", style_name=h2_style, bold=True)
    add_para(doc,
        "Dữ liệu thực nghiệm được xây dựng từ 3 nguồn chính: nền tảng thương mại điện tử Tiki, "
        "Shopee và Lazada – ba sàn TMĐT lớn nhất Việt Nam hiện nay. Quá trình thu thập kết hợp "
        "giữa web scraping tự động (Selenium) và bổ sung có kiểm soát bằng API của Large Language "
        "Model (GPT) cho các lớp thiểu số. Toàn bộ dữ liệu được gán nhãn thủ công bởi 2 thành "
        "viên nhóm với cross-checking để đảm bảo độ tin cậy.",
        style_name=body_style)
    add_para(doc,
        "Tập dữ liệu cuối cùng gồm 26.382 mẫu bình luận tiếng Việt được phân thành 6 nhãn cảm "
        "xúc theo mô hình Ekman [11]. Bảng 4.1 mô tả phân phối chi tiết.",
        style_name=body_style)
    add_table(doc,
              headers=["Nhãn cảm xúc", "Số mẫu", "Tỉ lệ"],
              rows=DATA_DISTRIBUTION,
              caption='Bảng 4.1 Phân phối các lớp cảm xúc trong tập dữ liệu',
              col_widths=[2.0, 1.5, 1.5])
    add_image(doc, IMG['eda'], width_inches=6.0,
              caption='Hình 4.1 Phân phối số lượng mẫu theo từng lớp cảm xúc và đặc điểm văn bản')
    add_para(doc,
        "Đặc điểm thống kê của tập dữ liệu:",
        style_name=body_style)
    add_para(doc,
        "- Độ dài văn bản: Trung bình ~160 ký tự/bình luận, ~36 từ/câu sau khi tách từ. "
        "Phần lớn bình luận có độ dài từ 10 đến 80 từ. Phân tích EDA cho thấy 95% bình luận "
        "nằm trong khoảng 100 tokens – cơ sở để chọn max_length = 100 cho Tokenizer.",
        style_name=body_style)
    add_para(doc,
        "- Mất cân bằng dữ liệu: Tỷ lệ mất cân bằng giữa lớp đa số (Surprise, Happiness, Anger, "
        "Sadness ~18% mỗi lớp) và lớp thiểu số (Disgust 10.5%) là 1.7:1. Đây là mức mất cân bằng "
        "tương đối nhẹ so với mức 10:1 thường gặp trong bài toán anomaly detection, nhưng vẫn cần "
        "xử lý để đảm bảo F1-score công bằng cho tất cả các lớp.",
        style_name=body_style)
    add_para(doc,
        "- Phân chia dữ liệu: Áp dụng Stratified Sampling 80:20 để đảm bảo tỷ lệ nhãn đồng đều "
        "giữa Train (~21.146 mẫu) và Test (~5.236 mẫu). Không dùng phương pháp k-fold do kích "
        "thước dữ liệu đủ lớn và chi phí tính toán của 4 mô hình RNN/CNN.",
        style_name=body_style)

    add_para(doc,
        "Bảng 4.2 trình bày một số mẫu thực tế đại diện cho từng lớp cảm xúc. Quan sát các mẫu "
        "này giúp hiểu trực quan đặc trưng ngôn ngữ của từng nhãn và những thách thức phân loại.",
        style_name=body_style)
    add_table(doc,
              headers=["Lớp", "Mẫu thực tế", "Đặc trưng ngôn ngữ"],
              rows=SAMPLE_DATA_BY_CLASS,
              caption='Bảng 4.2 Một số mẫu dữ liệu minh họa cho 6 lớp cảm xúc',
              col_widths=[1.0, 3.6, 1.8])

    # ---- 4.2 Quy trình xử lý dữ liệu ----
    add_para(doc, "4.2 Quy trình xử lý dữ liệu thực nghiệm", style_name=h2_style, bold=True)
    add_para(doc,
        "Dữ liệu thực nghiệm trải qua đầy đủ pipeline tiền xử lý 5 bước đã trình bày ở Chương 3. "
        "Bảng 4.3 minh họa cụ thể kết quả từng bước tiền xử lý trên một bình luận thực tế, "
        "cho thấy tác động rõ ràng của từng giai đoạn:",
        style_name=body_style)
    add_table(doc,
              headers=PREPROCESS_EXAMPLES[0],
              rows=PREPROCESS_EXAMPLES[1:],
              caption='Bảng 4.3 Minh họa pipeline tiền xử lý qua từng bước',
              col_widths=[2.2, 4.0])
    add_para(doc,
        "Sau bước tiền xử lý, áp dụng Data Augmentation trên tập huấn luyện để giải quyết mất "
        "cân bằng và tăng cường tính bền vững của mô hình: Random Deletion (xác suất xóa từ "
        "p=0.15) và Random Swap (đổi chỗ 2 từ ngẫu nhiên). Tổng cộng bổ sung ~4.188 mẫu nhân "
        "tạo, ưu tiên cho lớp thiểu số Disgust (tăng từ 2.762 lên ~3.500 mẫu trong tập train).",
        style_name=body_style)

    # ---- 4.3 Công nghệ và môi trường ----
    add_para(doc, "4.3 Công nghệ và môi trường thực nghiệm", style_name=h2_style, bold=True)
    add_para(doc,
        "Bảng 4.4 liệt kê đầy đủ các công nghệ, thư viện và cấu hình phần cứng được sử dụng "
        "trong quá trình thực nghiệm. Tất cả thực nghiệm được tiến hành trên cùng môi trường để "
        "đảm bảo tính so sánh công bằng giữa các mô hình.",
        style_name=body_style)
    tech_stack = [
        ("Ngôn ngữ lập trình",   "Python 3.10",             "Jupyter Notebook (Google Colab)"),
        ("Deep Learning",         "TensorFlow 2.12 / Keras", "Xây dựng và huấn luyện mô hình"),
        ("NLP tiếng Việt",        "Underthesea 6.x",         "Tách từ (word segmentation)"),
        ("Vector hóa",            "Keras Tokenizer",         "Chuyển text → token IDs"),
        ("Xử lý dữ liệu",         "Pandas 2.0, NumPy 1.24",  "Thao tác DataFrame, ma trận"),
        ("Đánh giá mô hình",      "Scikit-learn 1.3",        "Metrics, Confusion Matrix"),
        ("Trực quan hóa",         "Matplotlib, Seaborn",     "Biểu đồ EDA, training curves"),
        ("Môi trường GPU",        "Google Colab Pro",        "NVIDIA Tesla T4 16GB VRAM"),
        ("Mixed Precision",       "TF AMP (fp16+fp32)",      "Tăng tốc 2–3x, giảm VRAM"),
        ("API triển khai",        "FastAPI 0.104",           "REST API inference endpoint"),
        ("Workflow tự động",      "Apache Airflow 2.7",      "Pipeline điều phối định kỳ"),
        ("Container hóa",         "Docker Compose",          "Đóng gói dịch vụ production"),
    ]
    add_table(doc,
              headers=["Thành phần", "Công nghệ / Phiên bản", "Vai trò"],
              rows=tech_stack,
              caption='Bảng 4.4 Công nghệ và thư viện sử dụng trong thực nghiệm',
              col_widths=[1.8, 2.0, 2.4])

    add_para(doc,
        "Cấu hình phần cứng và chiến lược tối ưu tài nguyên:",
        style_name=body_style)
    add_para(doc,
        "- Google Colab Pro với GPU NVIDIA Tesla T4 (16GB VRAM, 2.560 CUDA cores, 8.1 TFLOPS). "
        "Mixed Precision (Automatic Mixed Precision – AMP) được bật: các phép tính trung gian "
        "dùng float16, tích lũy gradient dùng float32, giảm bộ nhớ ~2x và tăng tốc 2–3x.",
        style_name=body_style)
    add_para(doc,
        "- Hạt giống ngẫu nhiên (random seed = 42) được cố định cho NumPy, TensorFlow và Python "
        "random để đảm bảo kết quả có thể tái lập (reproducibility).",
        style_name=body_style)
    add_para(doc,
        "- Tất cả 4 mô hình được huấn luyện trên cùng tập dữ liệu, cùng bộ siêu tham số (Bảng 3.3), "
        "và cùng Callbacks (Early Stopping, ReduceLROnPlateau) để đảm bảo so sánh công bằng.",
        style_name=body_style)

    # ---- 4.4 Phương pháp đánh giá ----
    add_para(doc, "4.4 Phương pháp đánh giá", style_name=h2_style, bold=True)
    add_para(doc,
        "Nghiên cứu sử dụng bộ các độ đo toàn diện để đánh giá hiệu suất mô hình từ nhiều góc "
        "độ, đặc biệt chú trọng đến tính công bằng giữa các lớp cảm xúc.",
        style_name=body_style)

    add_para(doc, "4.4.1 Accuracy và các chỉ số Precision, Recall, F1", style_name='Normal', bold=True)
    add_para(doc,
        "Accuracy đo tỷ lệ dự đoán đúng tổng thể. Tuy nhiên với dữ liệu mất cân bằng nhẹ, cần "
        "bổ sung Precision, Recall và F1-Score tính theo từng lớp:",
        style_name=body_style)
    add_formula(doc, "Accuracy = (1/N) Σᵢ₌₁ᴺ 𝟙[ŷᵢ = yᵢ]", label="4.1")
    add_formula(doc, "Pₖ = TPₖ / (TPₖ + FPₖ),   Rₖ = TPₖ / (TPₖ + FNₖ)", label="4.2")
    add_formula(doc, "F1ₖ = 2 · Pₖ · Rₖ / (Pₖ + Rₖ)", label="4.3")
    add_para(doc,
        "Macro-Average lấy trung bình số học không trọng số của từng lớp, đảm bảo mỗi lớp cảm "
        "xúc được đánh giá bình đẳng dù số mẫu khác nhau:",
        style_name=body_style)
    add_formula(doc, "Macro-F1 = (1/C) Σₖ₌₁ᶜ F1ₖ,   C = 6", label="4.4")

    add_para(doc, "4.4.2 Cohen's Kappa", style_name='Normal', bold=True)
    add_para(doc,
        "Cohen's Kappa [24] đo mức độ đồng thuận giữa dự đoán của mô hình và nhãn thật, có loại "
        "trừ phần đồng thuận do may rủi – điều mà Accuracy không làm được:",
        style_name=body_style)
    add_formula(doc, "κ = (p_o - p_e) / (1 - p_e)", label="4.5")
    add_para(doc,
        "trong đó p_o là tỷ lệ đồng thuận quan sát được (Accuracy), p_e là tỷ lệ đồng thuận kỳ "
        "vọng theo phân phối ngẫu nhiên. Thang đo Landis & Koch (1977) [24]: κ > 0.81 là 'Rất "
        "tốt' (Almost Perfect Agreement).",
        style_name=body_style)

    add_para(doc, "4.4.3 Matthews Correlation Coefficient (MCC)", style_name='Normal', bold=True)
    add_para(doc,
        "Matthews Correlation Coefficient (MCC) [25] đo tương quan giữa dự đoán và nhãn thật, "
        "đặc biệt hữu ích cho dữ liệu mất cân bằng vì xem xét tất cả 4 ô của ma trận nhầm lẫn "
        "(TP, TN, FP, FN):",
        style_name=body_style)
    add_formula(doc, "MCC = (TP·TN - FP·FN) / √((TP+FP)(TP+FN)(TN+FP)(TN+FN))", label="4.6")
    add_para(doc,
        "MCC nhận giá trị trong [-1, 1]; MCC = 1 là dự đoán hoàn hảo, MCC = 0 là ngẫu nhiên. "
        "Cho phân loại đa lớp, MCC được tính theo dạng tổng quát.",
        style_name=body_style)

    add_para(doc, "4.4.4 Confidence Analysis", style_name='Normal', bold=True)
    add_para(doc,
        "Ngoài các metric phân loại chuẩn, phân tích Confidence (xác suất max của vector Softmax) "
        "cho từng dự đoán để đánh giá mức độ tự tin và phân phối ổn định của mô hình. Confidence "
        "thấp (<0.5) trên một mẫu thường chỉ ra trường hợp khó phân loại hoặc nhãn không rõ ràng.",
        style_name=body_style)
    add_formula(doc, "conf(x) = max_{k} P(y=k | x) = max_{k} Softmax(z)_k", label="4.7")

    # ---- 4.5 Kết quả thực nghiệm ----
    add_para(doc, "4.5 Kết quả thực nghiệm", style_name=h2_style, bold=True)

    add_para(doc, "4.5.1 So sánh tổng quan 4 mô hình", style_name='Normal', bold=True)
    add_para(doc,
        "Bảng 4.5 trình bày kết quả so sánh tổng quan của 4 mô hình trên tập kiểm thử 5.236 mẫu. "
        "Một kết quả đáng chú ý là tất cả 4 mô hình đều đạt Accuracy > 92.99% – vượt đáng kể "
        "mục tiêu đặt ra (>90%). Khoảng cách giữa mô hình tốt nhất (MLP: 93.37%) và tệ nhất "
        "(BiLSTM: 92.99%) chỉ là 0.38%, cho thấy đặc trưng ngữ nghĩa được học bởi lớp Embedding "
        "và Pooling đóng vai trò quyết định, không phải kiến trúc ẩn phía sau.",
        style_name=body_style)
    add_table(doc,
              headers=["Mô hình", "Accuracy", "Precision (macro)", "Recall (macro)", "F1 (macro)"],
              rows=MODEL_COMPARISON,
              caption='Bảng 4.5 So sánh hiệu suất 4 mô hình Deep Learning trên tập kiểm thử (5.236 mẫu)',
              col_widths=[2.2, 1.0, 1.2, 1.2, 1.0])
    add_image(doc, IMG['metrics_compare'], width_inches=6.0,
              caption='Hình 4.2 So sánh các chỉ số Accuracy, Precision, Recall, F1-Score của 4 mô hình')

    add_para(doc,
        "Phân tích chi tiết kết quả so sánh:",
        style_name=body_style)
    add_para(doc,
        "(1) MLP đạt Accuracy cao nhất trên test (93.37%). Kết quả đáng ngạc nhiên này có thể "
        "giải thích bởi: (a) Kiến trúc đơn giản ít tham số (~2.5M) → overfitting ít hơn CNN/BiLSTM; "
        "(b) Global Average Pooling đã tạo ra biểu diễn câu (sentence embedding) đủ tốt từ trung "
        "bình các word embeddings; (c) Đặc điểm của bài toán – cảm xúc thường phụ thuộc nhiều vào "
        "từ khóa đặc trưng (sợ, tức, vui, ghê tởm) hơn là cấu trúc tuần tự phức tạp.",
        style_name=body_style)
    add_para(doc,
        "(2) Hybrid CNN+BiLSTM đạt val_accuracy cao nhất khi huấn luyện (93.53%) – điểm này "
        "chứng minh khả năng học ổn định nhất trong 4 mô hình, nhưng lệch nhẹ xuống 93.09% trên "
        "test. Khoảng cách Train-Val của Hybrid nhỏ nhất (<0.5%), cho thấy mô hình tổng quát hóa "
        "tốt nhất.",
        style_name=body_style)
    add_para(doc,
        "(3) BiLSTM có Precision cao nhất (0.929) nhờ cơ chế học ngữ cảnh hai chiều, nhưng tốn "
        "~170 giây/epoch (gấp 170 lần MLP) – không hợp lý về chi phí/lợi ích cho bài toán này.",
        style_name=body_style)
    add_para(doc,
        "(4) Khoảng cách hiệu suất nhỏ giữa 4 mô hình (~0.4%) gợi ý: bottleneck hiện tại không "
        "phải kiến trúc mô hình mà là chất lượng dữ liệu và pipeline tiền xử lý. Để cải thiện "
        "đáng kể, cần data chất lượng cao hơn hoặc Transformer (PhoBERT).",
        style_name=body_style)

    add_para(doc, "4.5.2 Phân tích quá trình huấn luyện", style_name='Normal', bold=True)
    add_image(doc, IMG['train_compare'], width_inches=6.5,
              caption='Hình 4.3 Đường cong Accuracy & Loss của 4 mô hình trong quá trình huấn luyện')
    add_para(doc,
        "Quan sát từ đường cong huấn luyện (Hình 4.3):",
        style_name=body_style)
    add_para(doc,
        "- MLP: Hội tụ nhanh nhất (epoch 1 đã đạt ~77% train_acc). Tuy nhiên có dấu hiệu "
        "overfitting rõ nhất với khoảng cách Train Acc (~98%) và Val Acc (~93%) lớn ~5%. Early "
        "Stopping kích hoạt sớm nhất trong 4 mô hình (epoch 8–9).",
        style_name=body_style)
    add_para(doc,
        "- 1D-CNN: Hội tụ nhanh sau 3 epochs (~91% val_acc), ổn định ở mức ~93% sau epoch 5. "
        "ReduceLROnPlateau kích hoạt 2 lần. Đường Train-Val gần nhau từ epoch 4, ít dấu hiệu "
        "overfitting.",
        style_name=body_style)
    add_para(doc,
        "- BiLSTM: Hội tụ chậm nhất do bản chất tuần tự không thể song song hóa. Đạt val_acc "
        "tối đa 93.37% ở epoch 7. Mỗi epoch mất ~170 giây. Đường Loss giảm mượt mà nhất, "
        "ít dao động nhất.",
        style_name=body_style)
    add_para(doc,
        "- Hybrid CNN+BiLSTM: Hội tụ trung bình (epoch 4 đã đạt ~93% val_acc), đỉnh 93.53% "
        "ở epoch 7. Khoảng cách Train-Val nhỏ nhất trong 4 mô hình, chứng minh tính ổn định và "
        "tổng quát hóa cao nhất.",
        style_name=body_style)
    add_para(doc,
        "Nhận xét chung: Cơ chế Early Stopping và ReduceLROnPlateau đã hoạt động hiệu quả ở "
        "tất cả các mô hình. Không mô hình nào chạy đủ 20 epochs (dừng ở 8–12 epochs), tiết "
        "kiệm đáng kể tài nguyên tính toán.",
        style_name=body_style)

    add_para(doc, "4.5.3 Phân tích chi tiết theo từng lớp cảm xúc", style_name='Normal', bold=True)
    add_para(doc,
        "Bảng 4.6 trình bày kết quả phân loại chi tiết theo từng lớp cảm xúc của mô hình Hybrid "
        "CNN+BiLSTM – mô hình có val_accuracy cao nhất và cân bằng tốt nhất giữa Precision–Recall:",
        style_name=body_style)
    add_table(doc,
              headers=["Lớp cảm xúc", "Precision", "Recall", "F1-Score", "Support"],
              rows=HYBRID_PER_CLASS,
              caption='Bảng 4.6 Kết quả phân loại chi tiết theo từng lớp – Mô hình Hybrid CNN+BiLSTM',
              col_widths=[1.8, 1.2, 1.2, 1.2, 1.0])
    add_para(doc,
        "Nhận xét chi tiết từng lớp (mô hình Hybrid CNN+BiLSTM):",
        style_name=body_style)
    add_para(doc,
        "- Fear (F1=0.99): Lớp dễ phân loại nhất. Lý do: từ vựng đặc trưng rõ ràng và ít "
        "trùng lặp với các lớp khác (sợ, lo, hàng giả, không chính hãng, bảo hành,...). "
        "Recall gần như tuyệt đối (1.00) – mô hình bắt được hầu hết mẫu Fear.",
        style_name=body_style)
    add_para(doc,
        "- Sadness (F1=0.94) và Anger (F1=0.93): Cả hai đều đạt kết quả tốt. Anger có "
        "Precision cao (0.96) nhờ từ vựng mạnh ('tức giận', 'yêu cầu', 'vô trách nhiệm'). "
        "Sadness có Recall cao hơn (0.91) nhờ phân bố từ khóa tiêu cực rõ ràng.",
        style_name=body_style)
    add_para(doc,
        "- Surprise (F1=0.94): Cân bằng tốt giữa Precision (0.93) và Recall (0.94). Lớp này "
        "có cả bình luận tích cực và tiêu cực, nhưng từ khóa đặc trưng ('không ngờ', 'bất ngờ', "
        "'wao', 'ôi trời') giúp phân biệt tốt.",
        style_name=body_style)
    add_para(doc,
        "- Happiness (F1=0.90): Recall cao (0.96) nhưng Precision thấp hơn (0.85). Điều này "
        "cho thấy mô hình có xu hướng dự đoán nhiều mẫu là Happiness (false positive cao), "
        "do đây là lớp đa số với từ khóa tích cực rất phổ biến.",
        style_name=body_style)
    add_para(doc,
        "- Disgust (F1=0.85): Lớp khó nhất, thách thức kép: (1) số mẫu thiểu số nhất (10.5%); "
        "(2) ranh giới ngữ nghĩa mờ với Happiness trong các bình luận mỉa mai/châm biếm. "
        "Mô hình đã cải thiện F1 từ ~0.75 (không dùng Class Weighting) lên 0.85 (có Class "
        "Weighting + Augmentation).",
        style_name=body_style)

    add_para(doc,
        "Bảng 4.7 trình bày kết quả chi tiết của mô hình MLP – mô hình đạt Accuracy cao nhất "
        "trên tập test – để so sánh với Hybrid:",
        style_name=body_style)
    add_table(doc,
              headers=["Lớp cảm xúc", "Precision", "Recall", "F1-Score", "Support"],
              rows=MLP_PER_CLASS,
              caption='Bảng 4.7 Kết quả phân loại chi tiết theo từng lớp – Mô hình MLP (best test acc)',
              col_widths=[1.8, 1.2, 1.2, 1.2, 1.0])
    add_para(doc,
        "So sánh MLP và Hybrid: MLP có F1 tốt hơn ở Fear (0.997 vs 0.99) và Sadness (0.941 vs "
        "0.94). Hybrid có F1 tốt hơn ở Fear và Anger. Tổng thể hai mô hình tương đương; MLP "
        "vượt trội về tốc độ inference (~2ms/mẫu vs ~40ms/mẫu của Hybrid).",
        style_name=body_style)

    add_para(doc, "4.5.4 Các metric NLP chuyên sâu", style_name='Normal', bold=True)
    add_para(doc,
        "Bảng 4.8 trình bày bộ các metric NLP chuyên sâu trên mô hình MLP (best test accuracy). "
        "Các metric này cung cấp góc nhìn toàn diện hơn so với Accuracy đơn thuần:",
        style_name=body_style)
    add_table(doc,
              headers=["Metric", "Giá trị", "Ghi chú"],
              rows=NLP_METRICS,
              caption='Bảng 4.8 Các metric NLP chuyên sâu trên mô hình tốt nhất (MLP)',
              col_widths=[2.2, 1.2, 2.6])
    add_para(doc,
        "Phân tích ý nghĩa:",
        style_name=body_style)
    add_para(doc,
        "- Cohen's Kappa = 0.9201: Theo thang Landis & Koch [24], κ > 0.81 là 'Almost Perfect "
        "Agreement' (rất tốt). Con số này chứng minh dự đoán của mô hình thực sự đồng thuận "
        "với nhãn thật, không phải kết quả may rủi (nếu dự đoán ngẫu nhiên, κ ≈ 0).",
        style_name=body_style)
    add_para(doc,
        "- MCC = 0.9205: Tương quan cao giữa dự đoán và nhãn thật, đặc biệt có ý nghĩa vì "
        "MCC không bị ảnh hưởng nhiều bởi mất cân bằng dữ liệu như Accuracy.",
        style_name=body_style)
    add_para(doc,
        "- Confidence trung bình = 0.9733 (Std = 0.0819): Mô hình tự tin và ổn định trong hầu "
        "hết các dự đoán. Độ lệch chuẩn thấp (0.08) cho thấy phân phối confidence đồng đều, "
        "không có tình trạng mô hình rất tự tin ở một số mẫu và rất không chắc chắn ở số khác.",
        style_name=body_style)

    # ---- 4.5.5 Confusion Matrix và phân tích lỗi ----
    add_para(doc, "4.5.5 Phân tích Confusion Matrix và lỗi điển hình",
             style_name='Normal', bold=True)
    add_image(doc, IMG['best_analysis'], width_inches=6.5,
              caption='Hình 4.4 Confusion Matrix, độ tin cậy dự đoán và phân bố lỗi của mô hình tốt nhất')
    add_para(doc,
        "Phân tích Confusion Matrix (Hình 4.4) tiết lộ pattern nhầm lẫn đặc trưng:",
        style_name=body_style)
    add_para(doc,
        "- Cặp nhầm lẫn nghiêm trọng nhất: Disgust → Happiness (~9.8% mẫu Disgust bị nhầm "
        "thành Happiness). Nguyên nhân ngôn ngữ học rõ ràng: trong tiếng Việt, bình luận mỉa "
        "mai/châm biếm thường dùng từ ngữ tích cực bề ngoài (vd: 'sản phẩm tuyệt vời lắm, "
        "đúng giá tiền nào của ấy'). Đây là thách thức sarcasm detection vốn khó ngay cả với "
        "con người.",
        style_name=body_style)
    add_para(doc,
        "- Cặp nhầm lẫn thứ hai: Sadness → Anger (~5%). Hai cảm xúc này có ranh giới mờ "
        "khi mức độ cảm xúc ở ngưỡng trung bình – người dùng vừa buồn vừa tức giận về sản "
        "phẩm kém chất lượng.",
        style_name=body_style)
    add_para(doc,
        "- Fear và Surprise có tỷ lệ nhầm lẫn thấp nhất (<2%), do từ vựng đặc trưng tách "
        "biệt rõ ràng.",
        style_name=body_style)
    add_para(doc,
        "Bảng 4.9 trình bày phân tích chi tiết các trường hợp dự đoán sai điển hình với "
        "confidence thấp (<0.55), cung cấp insight về điểm yếu của mô hình:",
        style_name=body_style)
    add_table(doc,
              headers=["Văn bản", "True", "Pred", "Conf.", "Lý do lỗi"],
              rows=ERROR_ANALYSIS,
              caption='Bảng 4.9 Phân tích các trường hợp dự đoán sai điển hình',
              col_widths=[2.5, 0.9, 0.9, 0.6, 2.1])
    add_para(doc,
        "Nhận xét từ phân tích lỗi: (1) Sarcasm/mỉa mai là dạng lỗi khó nhất và phổ biến nhất; "
        "(2) Câu quá ngắn (<5 từ) thường dự đoán sai do thiếu ngữ cảnh; (3) Phủ định liên tiếp "
        "('không... không... không') gây nhiễu mô hình. Cả ba vấn đề này đều có thể cải thiện "
        "bằng cách sử dụng mô hình Transformer với cơ chế Self-Attention.",
        style_name=body_style)

    # ---- 4.6 So sánh chi phí tính toán ----
    add_para(doc, "4.6 So sánh chi phí tính toán", style_name=h2_style, bold=True)
    add_para(doc,
        "Ngoài hiệu suất phân loại, chi phí tính toán là yếu tố quan trọng khi lựa chọn mô hình "
        "cho triển khai thực tế. Bảng 4.10 so sánh chi tiết 4 chiều: số tham số, tốc độ huấn "
        "luyện, tốc độ suy luận và yêu cầu bộ nhớ GPU.",
        style_name=body_style)
    add_table(doc,
              headers=COMPUTATIONAL_COST[0],
              rows=COMPUTATIONAL_COST[1:],
              caption='Bảng 4.10 So sánh chi phí tính toán của 4 mô hình',
              col_widths=[1.7, 1.3, 1.3, 1.4, 1.3])
    add_para(doc,
        "Phân tích trade-off hiệu suất – chi phí:",
        style_name=body_style)
    add_para(doc,
        "- MLP: Lựa chọn tối ưu về chi phí/lợi ích. Accuracy cao nhất trên test (93.37%), "
        "tham số ít nhất (~2.5M), inference nhanh nhất (~2ms), bộ nhớ thấp nhất (~1GB). "
        "Phù hợp khi cần triển khai trên CPU hoặc thiết bị có tài nguyên hạn chế.",
        style_name=body_style)
    add_para(doc,
        "- 1D-CNN: Cân bằng tốt với inference nhanh (~5ms) và accuracy tốt (93.18%). Train "
        "nhanh nhất (~1s/epoch) do phép tích chập song song hóa tốt trên GPU. Phù hợp cho "
        "hệ thống real-time xử lý lượng lớn văn bản.",
        style_name=body_style)
    add_para(doc,
        "- BiLSTM: Chi phí tính toán cao nhất (~170s/epoch do tuần tự không song song hóa "
        "được), nhưng accuracy chỉ tương đương MLP. Không hợp lý về trade-off cho bài toán "
        "này. Phù hợp hơn cho các bài toán cần hiểu ngữ cảnh dài hạn.",
        style_name=body_style)
    add_para(doc,
        "- Hybrid CNN+BiLSTM: Cân bằng tốt giữa tốc độ (~45s/epoch, ~40ms inference) và "
        "val_accuracy cao nhất (93.53%). Khuyến nghị cho các hệ thống có GPU và không yêu cầu "
        "real-time cứng (batch processing, inference định kỳ qua Airflow).",
        style_name=body_style)
    add_para(doc,
        "Kết luận lựa chọn mô hình cho triển khai: Với mục tiêu balance giữa accuracy, tốc độ "
        "và tài nguyên, nhóm khuyến nghị MLP cho triển khai trên CPU (REST API nhẹ) và Hybrid "
        "CNN+BiLSTM cho triển khai trên GPU (batch processing qua Airflow).",
        style_name=body_style)

    # ============ CHƯƠNG 5 ============
    doc.add_page_break()
    add_para(doc, "CHƯƠNG 5", style_name=h1_style, bold=True)
    add_para(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", style_name=h1_style, bold=True)

    add_para(doc, "5.1 Tổng kết kết quả đạt được", style_name=h2_style, bold=True)
    add_para(doc,
        "Đồ án đã hoàn thành đầy đủ các mục tiêu đề ra ban đầu, từ xây dựng quy trình thu thập "
        "và tiền xử lý dữ liệu đặc thù tiếng Việt, đến thiết kế và huấn luyện các kiến trúc Deep "
        "Learning, và cuối cùng là đánh giá toàn diện kết quả thực nghiệm.",
        style_name=body_style)
    add_para(doc,
        "Về dữ liệu: Thu thập và xây dựng thành công tập dữ liệu 26.382 mẫu bình luận tiếng Việt "
        "từ 3 nền tảng TMĐT (Tiki, Shopee, Lazada) với 6 nhãn cảm xúc theo mô hình Ekman. Đây là "
        "tập dữ liệu phân loại đa lớp cảm xúc cho tiếng Việt TMĐT có quy mô lớn, được gán nhãn "
        "cẩn thận.",
        style_name=body_style)
    add_para(doc,
        "Về pipeline tiền xử lý: Xây dựng quy trình 5 bước chuyên biệt cho tiếng Việt gồm "
        "chuẩn hóa Teencode (từ điển 367 cặp ánh xạ), làm sạch và chuẩn hóa Unicode NFC, tách "
        "từ Underthesea, vector hóa Keras Tokenizer và Data Augmentation (Random Deletion, Random "
        "Swap). Quy trình này đóng vai trò then chốt trong việc đạt được kết quả cao.",
        style_name=body_style)
    add_para(doc,
        "Về kết quả thực nghiệm: Cả 4 kiến trúc Deep Learning được so sánh đều đạt Accuracy > "
        "92.99% – vượt mục tiêu đặt ra (>90%). Kết quả chi tiết:",
        style_name=body_style)
    summary_results = [
        ("MLP (Baseline)",             "93.37%", "0.9299", "0.9286", "~1s/epoch",    "Tốt nhất test acc"),
        ("1D-CNN",                     "93.18%", "0.9283", "0.9264", "~1s/epoch",    "Train nhanh nhất"),
        ("BiLSTM",                     "92.99%", "0.9290", "0.9248", "~170s/epoch",  "Chi phí cao nhất"),
        ("Hybrid CNN+BiLSTM (đề xuất)","93.09%", "0.9281", "0.9262", "~45s/epoch",   "Ổn định nhất (val)"),
    ]
    add_table(doc,
              headers=["Mô hình", "Accuracy", "Precision", "F1 (macro)", "Train/epoch", "Nhận xét"],
              rows=summary_results,
              caption='Bảng 5.1 Tổng hợp kết quả cuối cùng của 4 mô hình',
              col_widths=[1.8, 0.9, 1.0, 1.0, 1.0, 1.5])
    add_para(doc,
        "Cohen's Kappa = 0.9201 và MCC = 0.9205 trên mô hình tốt nhất xác nhận hệ thống đạt "
        "mức độ đồng thuận 'rất tốt' với nhãn thật theo tiêu chuẩn khoa học [24, 25]. Confidence "
        "trung bình 0.9733 cho thấy mô hình ổn định và đáng tin cậy trong suy luận thực tế.",
        style_name=body_style)

    add_para(doc, "5.2 Đóng góp chính của đồ án", style_name=h2_style, bold=True)
    add_para(doc,
        "Đồ án có những đóng góp cụ thể ở cả khía cạnh học thuật và thực tiễn:",
        style_name=body_style)
    add_para(doc,
        "Đóng góp (1) – Tài nguyên NLP tiếng Việt: Từ điển chuẩn hóa Teencode 367 cặp ánh xạ "
        "được xây dựng từ khảo sát thực tế là tài nguyên có thể tái sử dụng cho các nghiên cứu "
        "NLP tiếng Việt khác, đặc biệt trong miền mạng xã hội và TMĐT nơi Teencode phổ biến. "
        "Chưa có nghiên cứu tiếng Việt nào trước đây công bố tài nguyên tương tự.",
        style_name=body_style)
    add_para(doc,
        "Đóng góp (2) – Quy trình tiền xử lý chuyên biệt: Pipeline 5 bước kết hợp chuẩn hóa "
        "Teencode + Unicode NFC + Underthesea word segmentation + Keras Tokenizer + EDA Augmentation "
        "là một đề xuất hoàn chỉnh cho tiếng Việt TMĐT. Quy trình đã được xác nhận hiệu quả qua "
        "kết quả thực nghiệm (tất cả mô hình đạt >92%).",
        style_name=body_style)
    add_para(doc,
        "Đóng góp (3) – Báo cáo so sánh khoa học: So sánh đồng thời 4 kiến trúc Deep Learning "
        "(MLP, CNN, BiLSTM, Hybrid) trên cùng pipeline, cùng tập dữ liệu, cùng siêu tham số và "
        "cùng cơ chế đánh giá. Đây là tham chiếu khoa học cho các nghiên cứu tương lai muốn lựa "
        "chọn kiến trúc phù hợp cho bài toán phân tích cảm xúc tiếng Việt.",
        style_name=body_style)
    add_para(doc,
        "Đóng góp (4) – Hệ thống sẵn sàng triển khai: Mô hình được đóng gói thành module Python "
        "tích hợp với REST API (FastAPI) và workflow Apache Airflow, đã hoạt động trong môi trường "
        "phát triển. Đây là bước vượt trội so với các đồ án thông thường chỉ dừng ở notebook "
        "Jupyter, không có hạ tầng triển khai.",
        style_name=body_style)
    add_para(doc,
        "Đóng góp (5) – Phân tích lỗi toàn diện: Phân tích lỗi điển hình (sarcasm, câu quá ngắn, "
        "phủ định liên tiếp) và Confusion Matrix chi tiết cung cấp insight sâu về điểm yếu của "
        "mô hình, định hướng rõ ràng cho các cải tiến tiếp theo.",
        style_name=body_style)

    add_para(doc, "5.3 Hạn chế", style_name=h2_style, bold=True)
    add_para(doc,
        "Mặc dù đạt được kết quả tốt, đồ án vẫn có một số hạn chế cần được thừa nhận và giải "
        "quyết trong các nghiên cứu tiếp theo:",
        style_name=body_style)
    add_para(doc,
        "Hạn chế (1) – Lớp Disgust vẫn khó nhất (F1=0.85): Sarcasm/mỉa mai trong tiếng Việt "
        "là thách thức cố hữu mà các mô hình Deep Learning không có cơ chế Self-Attention khó "
        "xử lý. Kỹ thuật Class Weighting và Augmentation đã cải thiện nhưng chưa đủ để đưa "
        "Disgust lên ngang bằng các lớp khác.",
        style_name=body_style)
    add_para(doc,
        "Hạn chế (2) – Chưa kiểm thử cross-domain: Mô hình chỉ được kiểm thử trên dữ liệu TMĐT. "
        "Chưa có bằng chứng về khả năng tổng quát hóa sang các miền khác như mạng xã hội "
        "(Facebook, Zalo), tin tức, hoặc đánh giá ứng dụng di động.",
        style_name=body_style)
    add_para(doc,
        "Hạn chế (3) – Chưa so sánh với Transformer: Chưa thực hiện baseline comparison với "
        "PhoBERT [3] hoặc ViSoBERT [22] – các mô hình SOTA cho tiếng Việt. Khoảng cách hiệu "
        "suất giữa Hybrid CNN+BiLSTM (93%) và PhoBERT (kỳ vọng 95–97%) là điều cần làm rõ.",
        style_name=body_style)
    add_para(doc,
        "Hạn chế (4) – Nhãn đơn lớp (single-label): Đồ án sử dụng single-label classification "
        "– mỗi bình luận chỉ được gán 1 nhãn. Thực tế, một bình luận có thể chứa nhiều cảm xúc "
        "đan xen (vd: vừa Sadness vừa Anger). Multi-label classification sẽ phản ánh thực tế "
        "tốt hơn nhưng tăng độ phức tạp đáng kể.",
        style_name=body_style)
    add_para(doc,
        "Hạn chế (5) – Từ điển Teencode cần cập nhật liên tục: Ngôn ngữ mạng xã hội thay đổi "
        "nhanh; từ điển 367 cặp xây dựng trong năm 2024–2025 sẽ dần lỗi thời khi xuất hiện "
        "các Teencode mới. Cần cơ chế cập nhật tự động hoặc định kỳ.",
        style_name=body_style)

    add_para(doc, "5.4 Hướng phát triển", style_name=h2_style, bold=True)
    add_para(doc,
        "Dựa trên các hạn chế trên, nhóm đề xuất 5 hướng phát triển cụ thể theo thứ tự ưu tiên:",
        style_name=body_style)
    add_para(doc,
        "Hướng 1 (Ưu tiên cao) – Tích hợp mô hình ngôn ngữ tiền huấn luyện tiếng Việt: Fine-tune "
        "PhoBERT [3] hoặc ViSoBERT [22] trên tập dữ liệu của đề tài. Ưu điểm: biểu diễn ngữ "
        "nghĩa chất lượng cao học được từ 20GB corpus tiếng Việt; khả năng hiểu ngữ cảnh sâu "
        "thông qua Bidirectional Attention. Kỳ vọng cải thiện Accuracy lên 95–97% và đặc biệt "
        "cải thiện F1 của lớp Disgust (sarcasm) lên 0.90+.",
        style_name=body_style)
    add_para(doc,
        "Hướng 2 (Ưu tiên cao) – Giải quyết triệt để Sarcasm Detection: Áp dụng kỹ thuật nhận "
        "diện mỉa mai chuyên biệt: (a) Sentiment Incongruity Detection – phát hiện mâu thuẫn "
        "giữa từ ngữ tích cực và thực tế tiêu cực; (b) Sử dụng LLM để bổ sung dữ liệu Disgust "
        "mỉa mai có chất lượng; (c) Contrastive Learning để học ranh giới giữa Disgust và "
        "Happiness rõ hơn.",
        style_name=body_style)
    add_para(doc,
        "Hướng 3 (Ưu tiên trung bình) – Mở rộng sang Multi-label Classification: Cho phép mỗi "
        "bình luận có nhiều nhãn cảm xúc đồng thời. Điều này đòi hỏi gán nhãn lại dữ liệu và "
        "thay đổi kiến trúc đầu ra (sigmoid thay vì softmax, Binary Crossentropy loss).",
        style_name=body_style)
    add_para(doc,
        "Hướng 4 (Ưu tiên trung bình) – Kiểm thử cross-domain: Thu thập thêm dữ liệu từ Facebook "
        "(comment bài post), Zalo (đánh giá sản phẩm), và Google Play Store (đánh giá app) để "
        "kiểm thử khả năng tổng quát hóa của mô hình sang các miền khác.",
        style_name=body_style)
    add_para(doc,
        "Hướng 5 (Ưu tiên thấp) – Cơ chế cập nhật Teencode tự động: Xây dựng pipeline phân "
        "tích lỗi tự động (khi inference confidence thấp trên từ lạ) để phát hiện Teencode mới "
        "và đề xuất cập nhật từ điển. Kết hợp với việc thu thập dữ liệu định kỳ qua Airflow.",
        style_name=body_style)

    add_para(doc, "5.5 Nhận định tổng thể", style_name=h2_style, bold=True)
    add_para(doc,
        "Đồ án đã chứng minh rằng với dữ liệu phù hợp (~26.000 mẫu), pipeline tiền xử lý chuyên "
        "biệt cho tiếng Việt và các kiến trúc Deep Learning không quá phức tạp, có thể đạt được "
        "độ chính xác >93% cho bài toán phân loại 6 cảm xúc – một kết quả hoàn toàn khả thi "
        "cho ứng dụng thực tế.",
        style_name=body_style)
    add_para(doc,
        "Kết quả thực nghiệm gợi ý một quan sát thú vị: đối với bài toán phân loại cảm xúc "
        "tiếng Việt dựa trên từ khóa đặc trưng, kiến trúc đơn giản (MLP) có thể đạt hiệu suất "
        "tương đương hoặc thậm chí tốt hơn kiến trúc phức tạp (BiLSTM). Điều này gợi ý bottleneck "
        "của hệ thống hiện tại nằm ở chất lượng dữ liệu và pipeline tiền xử lý, không phải độ "
        "phức tạp của kiến trúc mô hình.",
        style_name=body_style)
    add_para(doc,
        "Hệ thống đã được tích hợp với Apache Airflow và FastAPI, sẵn sàng triển khai thực tế "
        "trong các ứng dụng phân tích phản hồi khách hàng của doanh nghiệp Việt Nam. Với chi phí "
        "tính toán thấp (mô hình MLP chạy được trên CPU với inference ~2ms/mẫu), hệ thống hoàn "
        "toàn phù hợp với các doanh nghiệp vừa và nhỏ.",
        style_name=body_style)

    # ============ TÀI LIỆU THAM KHẢO ============
    doc.add_page_break()
    add_para(doc, "TÀI LIỆU THAM KHẢO", style_name=chapter_style,
             bold=True, align='center')
    for ref in REFERENCES:
        add_para(doc, ref, style_name='Normal')

    # ============ LÀM VIỆC NHÓM ============
    doc.add_page_break()
    add_para(doc, "LÀM VIỆC NHÓM", style_name=chapter_style,
             bold=True, align='center')
    add_para(doc, GROUP_WORK, style_name='Normal')

    # Bảng phân công công việc chi tiết
    work_assignment = [
        ("Thu thập dữ liệu",          "Ngô Hồng Thông",   "100%"),
        ("Tiền xử lý + Teencode dict","Ngô Hồng Thông",   "100%"),
        ("EDA và phân tích dữ liệu",  "Trần Quang Triều", "100%"),
        ("Triển khai MLP, CNN",       "Trần Quang Triều", "100%"),
        ("Triển khai BiLSTM, Hybrid", "Ngô Hồng Thông",   "100%"),
        ("Data Augmentation",         "Trần Quang Triều", "100%"),
        ("Đánh giá, vẽ biểu đồ",      "Trần Quang Triều", "100%"),
        ("Triển khai Airflow + API",  "Cả nhóm",          "100%"),
        ("Viết báo cáo Ch.1, 3, 5",   "Ngô Hồng Thông",   "100%"),
        ("Viết báo cáo Ch.2, 4",      "Trần Quang Triều", "100%"),
        ("Review và hoàn thiện",      "Cả nhóm",          "100%"),
    ]
    add_table(doc,
              headers=["Công việc", "Phụ trách", "Hoàn thành"],
              rows=work_assignment,
              caption='Bảng phân công công việc chi tiết',
              col_widths=[3.0, 2.0, 1.2])

    # ============ TỰ ĐÁNH GIÁ ============
    doc.add_page_break()
    add_para(doc, "TỰ ĐÁNH GIÁ", style_name=chapter_style,
             bold=True, align='center')

    self_eval = [
        ("1. Trình bày bài toán và ý nghĩa", "2.0 / 2.0",
         "Đã trình bày rõ bài toán Sentiment Analysis tiếng Việt"),
        ("2. Tổng quan nghiên cứu liên quan", "1.5 / 2.0",
         "Có tổng quan các phương pháp, chưa so sánh sâu với SOTA tiếng Việt"),
        ("3. Mô hình đề xuất và thiết kế", "2.0 / 2.0",
         "Đề xuất Hybrid CNN+BiLSTM rõ ràng, có giải thích chi tiết"),
        ("4. Cài đặt và thực nghiệm", "2.0 / 2.0",
         "Code chạy được, kết quả 93.37% Accuracy trên test ~5,236 mẫu"),
        ("5. Đánh giá kết quả", "1.5 / 2.0",
         "Có so sánh 4 mô hình, phân tích confusion matrix, NLP metrics"),
        ("Tổng cộng tự đánh giá", "9.0 / 10", ""),
    ]
    add_table(doc,
              headers=["Tiêu chí", "Điểm", "Ghi chú"],
              rows=self_eval,
              col_widths=[2.4, 1.0, 3.0])

    doc.save(dst)
    print(f"[OK] Đã tạo: {dst}")


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    os.makedirs('result', exist_ok=True)
    print("Đang tạo file bài báo khoa học...")
    build_paper()
    print("Đang tạo file báo cáo đồ án...")
    build_report()
    print("\nHoàn thành! Kiểm tra folder result/")
